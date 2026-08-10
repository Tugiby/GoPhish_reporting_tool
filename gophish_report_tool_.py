import os
import sys
import traceback
import threading
import queue
import time
from dataclasses import dataclass
from datetime import datetime
from logging.handlers import RotatingFileHandler
from typing import Callable, Optional

# --- 1. KURŞUN GEÇİRMEZ HATA YAKALAYICI (CRASH REPORTER) ---
def global_exception_handler(exctype, value, tb):
    error_msg = "".join(traceback.format_exception(exctype, value, tb))
    
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
    except NameError:
        current_dir = os.getcwd()
        
    log_path = os.path.join(current_dir, "GOPHISH_HATA_LOG.txt")
    
    try:
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("=== GOPHISH ARACI HATA RAPORU ===\n")
            f.write(f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n\n")
            f.write(error_msg)
            f.write("\n=================================\n")
            f.write("Bu hatayı yapay zekaya veya geliştiriciye ileterek sorunu çözebilirsiniz.")
            
        if os.name == 'nt':
            os.startfile(log_path)
    except Exception:
        pass
    
    print(error_msg)
    print(f"\n[!] Kritik hata. Log dosyasi olusturuldu: {log_path}")
    try:
        input("Çıkmak için ENTER tuşuna basın...")
    except Exception:
        pass

sys.excepthook = global_exception_handler

# --- 2. OTOMATİK BAĞIMLILIK KONTROLÜ VE YÜKLEME ---
REQUIRED_PACKAGES = {
    "gophish": "gophish",
    "docx": "python-docx",
    "matplotlib": "matplotlib",
    "openpyxl": "openpyxl",
    "urllib3": "urllib3",
    "customtkinter": "customtkinter",
    "PIL": "Pillow",
}
if os.name == "nt":
    REQUIRED_PACKAGES["windnd"] = "windnd"

def auto_install_packages():
    import subprocess
    import importlib
    
    missing = []
    for import_name, pkg_name in REQUIRED_PACKAGES.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pkg_name)
            
    if missing:
        print(f"[*] Eksik paketler tespit edildi, otomatik yükleniyor: {', '.join(missing)}")
        subprocess.check_call([sys.executable, "-m", "pip", "install"] + missing)
        importlib.invalidate_caches()
        print("[+] Tüm bağımlılıklar başarıyla yüklendi!")

auto_install_packages()

# --- GEREKLİ KÜTÜPHANELERİN İÇE AKTARILMASI ---
import json
import io
import re
import math
import logging
import urllib3
import customtkinter as ctk
import tkinter as tk
import tkinter.font as tkfont
from tkinter import messagebox, filedialog

from gophish import Gophish
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from PIL import Image, ImageDraw

try:
    import windnd
except ImportError:
    windnd = None

# CustomTkinter Genel Ayarları
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

try:
    _current_dir = os.path.dirname(os.path.abspath(__file__))
except NameError:
    _current_dir = os.getcwd()
    
_log_file_path = os.path.join(_current_dir, 'gophish_report.log')

_log_handler = RotatingFileHandler(
    _log_file_path, maxBytes=5 * 1024 * 1024, backupCount=3, encoding="utf-8"
)
_log_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

logging.basicConfig(level=logging.INFO, handlers=[_log_handler])

try:
    urllib3.disable_warnings()
except AttributeError:
    pass

CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".gophish_config_v2.json")

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {
        "server": "",
        "api_key": "",
        "output_dir": os.path.join(os.path.expanduser("~"), "Desktop", "Gophish_Raporlari"),
        "mask_payload": True,
        "theme": "System",
        "lang": "TR",
        "word_enabled": True,
        "w_chart": True,
        "w_time": False,
        "w_payload": False,
        "w_browser": False,
        "w_dynamic_payload": False,
        "excel_enabled": True,
        "e_time": False,
        "e_payload": False,
        "e_browser": False,
        "e_dynamic_payload": False,
        "e_status": False,
        "e_domain_stats": False
    }

def save_config_all(config_data):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(config_data, f, indent=4)
    except Exception:
        pass


# --- FAZ 1: ALTYAPI (Threading, Cancel, Timeout/Retry, Toast, Progress) ---

API_TIMEOUT_SEC = 30
API_MAX_RETRIES = 3
API_RETRY_BASE_DELAY = 1.5


class CancelledError(Exception):
    """İşlem kullanıcı tarafından iptal edildi."""


@dataclass
class ReportOptions:
    """Rapor üretim seçenekleri."""
    mask_payload: bool = True
    word_enabled: bool = True
    w_chart: bool = True
    w_time: bool = False
    w_payload: bool = False
    w_browser: bool = False
    w_dynamic_payload: bool = False
    excel_enabled: bool = True
    e_time: bool = False
    e_payload: bool = False
    e_browser: bool = False
    e_dynamic_payload: bool = False
    e_status: bool = False
    e_domain_stats: bool = False

    @classmethod
    def from_dict(cls, data: dict) -> "ReportOptions":
        return cls(**{k: data[k] for k in cls.__dataclass_fields__ if k in data})


def count_report_steps(camp_ids: list, options: dict) -> int:
    """Rapor pipeline'ındaki toplam adım sayısını hesaplar."""
    steps = 1  # API'den veri çekme
    steps += 1  # Veri işleme / analiz
    if options.get("word_enabled"):
        steps += 1
    if options.get("excel_enabled"):
        steps += 1
    return max(steps, 1)


class ProgressReporter:
    """Adımlı ilerleme takibi (1/N formatında)."""

    def __init__(self, total_steps: int, callback: Callable[[int, int, str, float], None]):
        self.total_steps = max(total_steps, 1)
        self.current_step = 0
        self.callback = callback

    def advance(self, message: str, cancel_event: Optional[threading.Event] = None) -> None:
        if cancel_event and cancel_event.is_set():
            raise CancelledError()
        self.current_step += 1
        percent = (self.current_step / self.total_steps) * 100
        self.callback(self.current_step, self.total_steps, message, percent)

    @staticmethod
    def check_cancel(cancel_event: Optional[threading.Event]) -> None:
        if cancel_event and cancel_event.is_set():
            raise CancelledError()


class GophishApiClient:
    """Timeout ve retry destekli GoPhish API sarmalayıcısı."""

    def __init__(
        self,
        api_key: str,
        server: str,
        cancel_event: Optional[threading.Event] = None,
        timeout: int = API_TIMEOUT_SEC,
        max_retries: int = API_MAX_RETRIES,
    ):
        self.api_key = api_key
        self.server = server
        self.cancel_event = cancel_event
        self.timeout = timeout
        self.max_retries = max_retries
        self._api: Optional[Gophish] = None

    def _get_api(self) -> Gophish:
        if self._api is None:
            self._api = Gophish(self.api_key, host=self.server, verify=False)
        return self._api

    def _run_with_retry(self, operation: Callable, operation_name: str):
        last_error: Optional[Exception] = None

        for attempt in range(1, self.max_retries + 1):
            ProgressReporter.check_cancel(self.cancel_event)

            result_box: list = []
            error_box: list = []

            def _target():
                try:
                    result_box.append(operation())
                except Exception as exc:
                    error_box.append(exc)

            worker = threading.Thread(target=_target, daemon=True)
            worker.start()
            worker.join(timeout=self.timeout)

            if worker.is_alive():
                last_error = TimeoutError(
                    f"{operation_name} zaman aşımına uğradı ({self.timeout}s)"
                )
                logging.warning(
                    "%s deneme %d/%d: timeout", operation_name, attempt, self.max_retries
                )
            elif error_box:
                last_error = error_box[0]
                logging.warning(
                    "%s deneme %d/%d: %s", operation_name, attempt, self.max_retries, last_error
                )
            else:
                return result_box[0]

            if attempt < self.max_retries:
                delay = API_RETRY_BASE_DELAY * (2 ** (attempt - 1))
                time.sleep(delay)

        raise last_error or RuntimeError(f"{operation_name} başarısız oldu.")

    def get_campaigns(self):
        return self._run_with_retry(
            lambda: self._get_api().campaigns.get(),
            "Kampanya listesi",
        )

    def get_campaign(self, campaign_id: int):
        return self._run_with_retry(
            lambda: self._get_api().campaigns.get(campaign_id=int(campaign_id)),
            f"Kampanya #{campaign_id}",
        )


class ToastManager:
    """Sağ altta kayarak beliren bildirimler."""

    COLORS = {
        "info": ("#2980b9", "#ffffff"),
        "success": ("#27ae60", "#ffffff"),
        "warning": ("#f39c12", "#ffffff"),
        "error": ("#c0392b", "#ffffff"),
    }

    def __init__(self, root: ctk.CTk):
        self.root = root
        self._container = ctk.CTkFrame(root, fg_color="transparent")
        self._container.place(relx=1.0, rely=1.0, anchor="se", x=-20, y=-20)
        self._toasts: list[ctk.CTkFrame] = []

    def show(self, message: str, level: str = "info", duration_ms: int = 4500) -> None:
        bg, fg = self.COLORS.get(level, self.COLORS["info"])

        toast = ctk.CTkFrame(self._container, fg_color=bg, corner_radius=8)
        toast.pack(pady=4, anchor="e")

        lbl = ctk.CTkLabel(
            toast, text=message, text_color=fg,
            font=ctk.CTkFont(size=12), wraplength=340, justify="left",
        )
        lbl.pack(padx=14, pady=10)

        self._toasts.append(toast)
        self._reposition_toasts()

        def dismiss():
            if toast.winfo_exists():
                toast.destroy()
                if toast in self._toasts:
                    self._toasts.remove(toast)
                self._reposition_toasts()

        self.root.after(duration_ms, dismiss)

    def _reposition_toasts(self) -> None:
        pass  # pack anchor="e" yeterli; üst üste dizilir


# --- FAZ 2: UX (İkonlar, Font, Drag&Drop, Spinner, Hover) ---

FONT_CANDIDATES = ("Inter", "SF Pro Display", "Roboto", "Segoe UI", "Helvetica Neue", "Arial")
ICONS_DIR = os.path.join(_current_dir, "assets", "icons")


class AppFonts:
    """Sistemde mevcut modern font ailesini seçer ve önbelleğe alır."""

    _family: Optional[str] = None
    _cache: dict = {}

    @classmethod
    def family(cls) -> str:
        if cls._family is None:
            try:
                available = set(tkfont.families())
            except Exception:
                available = set()
            cls._family = next((f for f in FONT_CANDIDATES if f in available), "Segoe UI")
        return cls._family

    @classmethod
    def get(cls, size: int = 13, weight: str = "normal", slant: str = "roman") -> ctk.CTkFont:
        key = (size, weight, slant)
        if key not in cls._cache:
            cls._cache[key] = ctk.CTkFont(family=cls.family(), size=size, weight=weight, slant=slant)
        return cls._cache[key]


class IconManager:
    """Pillow ile programatik ikon üretir ve CTkImage olarak sunar."""

    _ctk_cache: dict = {}

    PALETTE = {
        "shield": ("#1f6aa5", "#ffffff"),
        "search": ("#7f8c8d", "#ffffff"),
        "refresh": ("#2980b9", "#ffffff"),
        "folder": ("#f39c12", "#ffffff"),
        "sync": ("#27ae60", "#ffffff"),
        "campaign": ("#8e44ad", "#ffffff"),
        "report": ("#16a085", "#ffffff"),
        "settings": ("#566573", "#ffffff"),
        "back": ("#566573", "#ffffff"),
    }

    @classmethod
    def _ensure_dir(cls) -> None:
        os.makedirs(ICONS_DIR, exist_ok=True)

    @classmethod
    def _draw(cls, name: str, size: int) -> Image.Image:
        bg, fg = cls.PALETTE.get(name, ("#34495e", "#ffffff"))
        img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        m = max(2, size // 8)
        draw.rounded_rectangle([0, 0, size - 1, size - 1], radius=size // 5, fill=bg)

        cx, cy = size // 2, size // 2
        lw = max(1, size // 10)

        if name == "shield":
            draw.polygon([(cx, m + 1), (size - m - 1, m + 3), (size - m - 2, cy + 1), (cx, size - m - 1), (m + 1, cy + 1)], fill=fg)
        elif name == "search":
            r = size // 4
            draw.ellipse([cx - r, cy - r - 1, cx + r - 1, cy + r - 2], outline=fg, width=lw)
            draw.line([cx + r - 2, cy + r - 3, size - m - 2, size - m - 2], fill=fg, width=lw)
        elif name == "refresh":
            draw.arc([m + 1, m + 1, size - m - 1, size - m - 1], start=200, end=340, fill=fg, width=lw)
            draw.polygon([(size - m - 3, m + 2), (size - m - 1, m + 6), (size - m - 7, m + 5)], fill=fg)
        elif name == "folder":
            draw.rectangle([m + 1, cy - 1, size - m - 1, size - m - 1], fill=fg)
            draw.rectangle([m + 1, cy - 4, cx + 2, cy - 1], fill=fg)
        elif name == "sync":
            draw.line([m + 2, cy, size - m - 2, cy], fill=fg, width=lw)
            draw.polygon([(size - m - 3, cy - 3), (size - m - 1, cy), (size - m - 3, cy + 3)], fill=fg)
            draw.polygon([(m + 2, cy - 3), (m, cy), (m + 2, cy + 3)], fill=fg)
        elif name == "campaign":
            draw.rectangle([m + 2, m + 3, size - m - 2, size - m - 2], outline=fg, width=lw)
            draw.line([m + 4, cy, size - m - 4, cy], fill=fg, width=lw)
        elif name == "report":
            draw.rectangle([m + 3, m + 2, size - m - 3, size - m - 2], outline=fg, width=lw)
            draw.line([m + 5, cy - 2, size - m - 5, cy - 2], fill=fg, width=lw)
            draw.line([m + 5, cy + 2, size - m - 5, cy + 2], fill=fg, width=lw)
        elif name == "settings":
            r = size // 5
            draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill=fg)
            for angle in range(0, 360, 45):
                rad = math.radians(angle)
                x1 = cx + int((r + 2) * math.cos(rad))
                y1 = cy + int((r + 2) * math.sin(rad))
                x2 = cx + int((r + 5) * math.cos(rad))
                y2 = cy + int((r + 5) * math.sin(rad))
                draw.line([x1, y1, x2, y2], fill=fg, width=lw)
        elif name == "back":
            draw.polygon([(cx + 3, m + 2), (m + 2, cy), (cx + 3, size - m - 2)], fill=fg)
            draw.line([m + 2, cy, size - m - 2, cy], fill=fg, width=lw)
        else:
            draw.ellipse([m + 2, m + 2, size - m - 2, size - m - 2], fill=fg)

        return img

    @classmethod
    def get(cls, name: str, size: int = 20) -> Optional[ctk.CTkImage]:
        key = (name, size)
        if key in cls._ctk_cache:
            return cls._ctk_cache[key]

        cls._ensure_dir()
        path = os.path.join(ICONS_DIR, f"{name}_{size}.png")
        if not os.path.exists(path):
            cls._draw(name, size * 2).save(path)  # retina için 2x kaydet

        pil_img = Image.open(path).convert("RGBA")
        img = ctk.CTkImage(light_image=pil_img, dark_image=pil_img, size=(size, size))
        cls._ctk_cache[key] = img
        return img


class LoadingSpinner:
    """Buton üzerinde dönen yükleme animasyonu."""

    FRAMES = ("⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏")

    def __init__(self, widget, root: ctk.CTk):
        self.widget = widget
        self.root = root
        self._running = False
        self._idx = 0
        self._prefix = ""
        self._original_text = ""

    def start(self, prefix: str, original_text: str = "") -> None:
        self._running = True
        self._prefix = prefix
        self._original_text = original_text or getattr(self.widget, "cget", lambda x: "")("text")
        self._idx = 0
        self._tick()

    def stop(self, text: Optional[str] = None) -> None:
        self._running = False
        if text is not None:
            self.widget.configure(text=text)

    def _tick(self) -> None:
        if not self._running:
            return
        self.widget.configure(text=f"{self.FRAMES[self._idx]} {self._prefix}")
        self._idx = (self._idx + 1) % len(self.FRAMES)
        self.root.after(90, self._tick)


def setup_folder_drop(widget, callback: Callable[[str], None]) -> None:
    """Windows'ta klasör sürükle-bırak desteği ekler."""
    if windnd is None:
        return

    target = widget._entry if hasattr(widget, "_entry") else widget

    def on_drop(files):
        if not files:
            return
        raw = files[0]
        path = raw.decode("utf-8") if isinstance(raw, bytes) else str(raw)
        path = path.strip().strip("{}")
        if os.path.isdir(path):
            callback(path)
        elif os.path.isfile(path):
            callback(os.path.dirname(path))

    try:
        windnd.hook_dropfiles(target, func=on_drop)
    except Exception as exc:
        logging.debug("Drag-drop kurulamadı: %s", exc)


# --- FAZ 3: AKILLI OZELLIKLER (Cache, Risk Score, API Health, Dashboard) ---

CACHE_DIR = os.path.join(_current_dir, "cache")
CACHE_TTL_SEC = 300  # 5 dakika


class AppCache:
    """Kampanya listesi/detayi icin JSON tabanli TTL cache."""

    def __init__(self, ttl: int = CACHE_TTL_SEC):
        self.ttl = ttl
        os.makedirs(CACHE_DIR, exist_ok=True)

    def _path(self, key: str) -> str:
        # Guvenli dosya adi
        safe = re.sub(r"[^A-Za-z0-9_.-]", "_", str(key))
        return os.path.join(CACHE_DIR, f"{safe}.json")

    def get(self, key: str):
        """Cache'ten veriyi dondur; yoksa veya TTL bitmisse None."""
        path = self._path(key)
        try:
            if not os.path.exists(path):
                return None
            age = time.time() - os.path.getmtime(path)
            if age > self.ttl:
                return None
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return None

    def set(self, key: str, data) -> None:
        """Veriyi cache'e kaydet."""
        try:
            with open(self._path(key), "w", encoding="utf-8") as f:
                json.dump(data, f, default=str, ensure_ascii=False, indent=2)
        except Exception as exc:
            logging.debug("Cache yazma hatasi: %s", exc)

    def clear_all(self) -> None:
        try:
            for fname in os.listdir(CACHE_DIR):
                os.remove(os.path.join(CACHE_DIR, fname))
        except Exception:
            pass


def compute_risk_score(opened: int, clicked: int, submitted: int, total: int) -> float:
    """Agirlikli risk skoru: submit*3 + click*2 + open*1, normalize edilmis (0-100)."""
    if total <= 0:
        return 0.0
    raw = (submitted * 3 + clicked * 2 + opened * 1) / (total * 3) * 100
    return round(min(raw, 100.0), 1)


def risk_level(score: float) -> tuple[str, str]:
    """Risk seviyesi ve renk (TR label, hex renk)."""
    if score >= 40:
        return "Kritik", "#e74c3c"
    if score >= 20:
        return "Orta", "#f39c12"
    return "Dusuk", "#27ae60"


class ApiHealth:
    """API baglanti durumu ve gecikme olcumu."""

    def __init__(self, cancel_event=None, timeout: int = 10):
        self.cancel_event = cancel_event
        self.timeout = timeout
        self.status: str = "unknown"   # unknown / ok / slow / error
        self.latency_ms: Optional[float] = None
        self.checked_at: Optional[str] = None
        self.last_error: Optional[str] = None

    def check(self, server: str, api_key: str) -> dict:
        """Ping endpoint'e istek atar; durum + latency dondurur."""
        result = {"status": self.status, "latency_ms": self.latency_ms, "error": self.last_error}
        if not server or not api_key:
            self.status = "error"
            self.last_error = "Bilgi eksik"
            result = {"status": self.status, "latency_ms": None, "error": self.last_error}
            return result

        done = threading.Event()
        out = {}

        def _probe():
            try:
                start = time.time()
                import urllib.request
                req = urllib.request.Request(server.rstrip("/") + "/api/campaigns/")
                req.add_header("Authorization", api_key)
                with urllib.request.urlopen(req, timeout=self.timeout, context=__import__("ssl")._create_unverified_context()) as resp:
                    out["latency_ms"] = round((time.time() - start) * 1000, 1)
                    out["http"] = resp.status
            except Exception as exc:
                out["error"] = str(exc)

        worker = threading.Thread(target=_probe, daemon=True)
        worker.start()
        worker.join(timeout=self.timeout)

        if worker.is_alive():
            self.status = "error"
            self.latency_ms = None
            self.last_error = "Zaman asimi"
        else:
            self.latency_ms = out.get("latency_ms")
            self.last_error = out.get("error")
            if out.get("error"):
                self.status = "error"
            elif self.latency_ms and self.latency_ms > 1000:
                self.status = "slow"
            else:
                self.status = "ok"

        self.checked_at = datetime.now().strftime("%H:%M:%S")
        return {
            "status": self.status,
            "latency_ms": self.latency_ms,
            "error": self.last_error,
            "checked_at": self.checked_at,
        }


def compute_campaign_stats(campaign) -> dict:
    """Kampanya metriklerini bir kez hesaplar (UI performansı için)."""
    results = campaign.results or []
    total = len(results)
    sub = sum(1 for r in results if r.status == "Submitted Data")
    clk = sum(1 for r in results if r.status in ("Clicked Link", "Submitted Data"))
    opn = sum(1 for r in results if r.status in ("Email Opened", "Clicked Link", "Submitted Data"))
    return {"total": total, "opened": opn, "clicked": clk, "submitted": sub}


# --- ÇOKLU DİL SÖZLÜĞÜ (TR / EN) ---
TRANSLATIONS = {
    "TR": {
        "app_title": "GoPhish Profesyonel Raporlama Motoru",
        "nav_campaigns": "Kampanyalar",
        "nav_reports": "Rapor Tasarımı",
        "nav_settings": "Ayarlar",
        "server_url": "Sunucu URL:",
        "api_key": "API Key:",
        "sync_btn": "Kampanyaları Senkronize Et",
        "search_lbl": "Ara:",
        "completed_chk": "Sadece Tamamlananlar",
        "date_lbl": "Tarih:",
        "refresh_btn": "Yenile",
        "select_camp_msg": "Listeden en az bir kampanya seçin...",
        "target_card": "Hedeflenen",
        "open_card": "Açanlar",
        "click_card": "Tıklayanlar",
        "submit_card": "Veri Girenler",
        "output_dir_lbl": "Çıktı Klasörü:",
        "browse_btn": "Klasör Seç",
        "drop_hint": "Klasörü buraya sürükleyip bırakın",
        "opsec_frame": " Güvenlik ve Gizlilik (OPSEC) ",
        "mask_chk": "Hassas Verileri (Payload) Gizle (Parolaları *** yap)",
        "word_frame": " Word Raporu (.docx) ",
        "excel_frame": " Excel Raporu (.xlsx) ",
        "create_report_btn": "Seçili Kampanyaların Raporunu Oluştur",
        "settings_title": "Uygulama ve Görünüm Ayarları",
        "theme_lbl": "Tema Modu:",
        "lang_lbl": "Dil / Language:",
        "proceed_btn": "Seçilenlerle Rapor Ekranına İlerle ➔",
        "back_btn": " ⬅ Geri Dön",
        "cancel_btn": "İptal Et",
        "step_fetch": "Veriler çekiliyor...",
        "step_process": "Veriler işleniyor...",
        "step_word": "Word raporu oluşturuluyor...",
        "step_excel": "Excel raporu oluşturuluyor...",
        "step_done": "İşlem tamamlandı!",
        "step_cancelled": "İşlem iptal edildi.",
        "toast_sync_ok": "{count} kampanya senkronize edildi.",
        "toast_report_ok": "Raporlar oluşturuldu ({count} dosya).",
        "toast_conn_err": "Bağlantı hatası: {msg}",
        "toast_perm_err": "Dosya erişim hatası: {path}",
        "warn_credentials": "Sunucu URL ve API Key zorunludur.",
        "warn_no_campaign": "Lütfen en az bir kampanya seçin.",
        "warn_no_format": "En az bir çıktı formatı seçmelisiniz.",
        "generating": "Raporlar üretiliyor...",
        "connecting": "Bağlanıyor...",
        "api_health": "API Sagligi:",
        "not_checked": "Hic kontrol edilmedi",
        "latency": "Gecikme:",
        "last_sync": "Son Senkronizasyon:",
        "check_btn": "Kontrol Et",
        "trend_title": "Kampanya Trend Analizi  (Grafik)",
        "refresh_dash": "Dashboard Yenile",
        "viewing_camps": "Goruntulenen Kampanyalar:",
        "active_select": "Aktif Secimler",
        "all_campaigns": "Tum Kampanyalar",
        "card_total": "Toplam Kampanya",
        "card_target": "Toplam Hedef",
        "card_opened": "Acan Kisi",
        "card_clicked": "Tiklayan Kisi",
        "card_submitted": "Veri Giren",
        "card_risk": "Risk Skoru",
        "risk_critical": "Kritik",
        "risk_medium": "Orta",
        "risk_low": "Dusuk",
        "health_ok": "Bagli (OK)",
        "health_slow": "Yavas",
        "health_err": "Baglanti Hatasi",
        "no_chart": "Henuz grafik yok.\nKampanya senkronize edin veya secim yapin.",
        "dash_no_selection_msg": "Trend icin en az bir kampanya secin.",
        "dash_empty": "Depoda kampanya yok.\nOnce 'Kampanyalar' sekmesinden senkronize edin.",
        "no_campaigns": "Gosterilecek kampanya bulunamadi.",
        "report_title_multi": "Coklu Kampanya Trend Analizi Raporu",
        "report_desc_multi": "Bu rapor {n} farkli kampanyanin kiyaslamasini icermektedir.",
        "report_click_pct": "Tiklama (%)",
        "report_submit_pct": "Veri Girme (%)",
        "report_risk_trend_title": "Zaman Icindeki Risk Egilimi",
        "report_oran": "Oran (%)",
        "trend_h_campaign": "Kampanya",
        "trend_h_date": "Tarih",
        "trend_h_target": "Hedef",
        "trend_h_clicked": "Tiklayan (%)",
        "trend_h_submitted": "Veri Giren (%)",
        "report_title_single": "Kurumsal Oltalama Simulasyon Raporu",
        "report_campaign_lbl": "Kampanya: {name}",
        "report_exec_summary": "1. Yonetici Ozeti",
        "report_pie_critical": "Kritik (Veri)",
        "report_pie_medium": "Orta (Tik)",
        "report_pie_low": "Dusuk (Acan)",
        "report_pie_ignore": "Ilgilenmeyen",
        "report_stat_target": "Hedeflenen Personel",
        "report_stat_opened": "Acanlarin Toplami",
        "report_stat_clicked": "Tiklayanlarin Toplami",
        "report_stat_submitted": "Veri Girenler (Kritik)",
        "report_stat_leak": "Zafiyet (Sizinti) Orani",
        "report_detail_findings": "2. Detayli Bulgular",
        "report_table_critical": "Kritik Risk: Veri Giren Kullanicilar",
        "report_table_medium": "Orta Risk: Linke Tiklayan Kullanicilar",
        "table_name_col": "Isim Soyisim",
        "table_email_col": "E-posta Adresi",
        "table_time_col": "Eylem Saati",
        "table_payload_col": "Girilen Veri (Payload)",
        "table_device_col": "Cihaz",
        "table_browser_col": "Tarayıcı",
        "table_empty": "Bu kategoride herhangi bir kullanici bulunmamaktadir.",
        "payload_form_empty": "Form Bos",
        "payload_unreadable": "Veri okunamadi",
        "payload_masked": "*** (Gizlendi)",
        "browser_unavailable": "Cihaz Bilgisi Alinamadi",
        "excel_title_trend": "Trend Kiyaslamasi",
        "excel_title_campaign": "Kampanya Raporu",
        "file_report": "_Rapor",
        "file_trend": "Trend_Analizi",
        "time_all": "Tum Zamanlar",
        "time_1d": "Son 1 Gun",
        "time_5d": "Son 5 Gun",
        "time_1w": "Son 1 Hafta",
        "time_1m": "Son 1 Ay",
        "selected_msg": "{n} Kampanya Secildi",
        "w_generate": "Word Raporu Olustur",
        "w_add_chart": "Grafik Analizi Ekle",
        "w_add_time": "Eylem Saatlerini Ekle",
        "w_add_payload": "Veri (Payload) Ekle",
        "w_add_browser": "Cihaz/Tarayici Ekle",
        "w_add_dynamic": "Form Alanlarini Ayri Sutun Olarak Ekle",
        "e_generate": "Excel Raporu Olustur",
        "e_add_time": "Eylem Saatlerini Ekle",
        "e_add_payload": "Girilen Verileri (Payload) Ekle",
        "e_add_browser": "Cihaz/Tarayici Bilgisi Ekle",
        "e_add_dynamic": "Form Alanlarini Ayri Sutun Olarak Ekle",
        "e_add_status": "Veri Girdisinde Saglik Durumu (Geçerli/Şüpheli/False Positive) Ekle",
        "e_add_domain_stats": "Domain Bazlı Dağılım Sayfası Ekle",
    },
    "EN": {
        "app_title": "GoPhish Professional Reporting Engine",
        "nav_campaigns": "Campaigns",
        "nav_reports": "Report Design",
        "nav_settings": "Settings",
        "server_url": "Server URL:",
        "api_key": "API Key:",
        "sync_btn": "Synchronize Campaigns",
        "search_lbl": "Search:",
        "completed_chk": "Completed Only",
        "date_lbl": "Date:",
        "refresh_btn": "Refresh",
        "select_camp_msg": "Select at least one campaign...",
        "target_card": "Targeted",
        "open_card": "Opened",
        "click_card": "Clicked",
        "submit_card": "Submitted",
        "output_dir_lbl": "Output Directory:",
        "browse_btn": "Browse",
        "opsec_frame": " Security & Privacy (OPSEC) ",
        "mask_chk": "Mask Sensitive Payload Data (Replace passwords with ***)",
        "word_frame": " Word Report (.docx) ",
        "excel_frame": " Excel Report (.xlsx) ",
        "create_report_btn": "Generate Reports for Selected Campaigns",
        "settings_title": "Application & Appearance Settings",
        "theme_lbl": "Theme Mode:",
        "lang_lbl": "Language / Dil:",
        "proceed_btn": "Proceed to Reports ➔",
        "back_btn": " ⬅ Back",
        "cancel_btn": "Cancel",
        "step_fetch": "Fetching data...",
        "step_process": "Processing data...",
        "step_word": "Generating Word report...",
        "step_excel": "Generating Excel report...",
        "step_done": "Completed!",
        "step_cancelled": "Operation cancelled.",
        "toast_sync_ok": "{count} campaigns synchronized.",
        "toast_report_ok": "Reports created ({count} files).",
        "toast_conn_err": "Connection error: {msg}",
        "toast_perm_err": "File access error: {path}",
        "warn_credentials": "Server URL and API Key are required.",
        "warn_no_campaign": "Please select at least one campaign.",
        "warn_no_format": "Select at least one output format.",
        "generating": "Generating reports...",
        "connecting": "Connecting...",
        "api_health": "API Health:",
        "not_checked": "Never checked",
        "latency": "Latency:",
        "last_sync": "Last Sync:",
        "check_btn": "Check",
        "trend_title": "Campaign Trend Analysis  (Chart)",
        "refresh_dash": "Refresh Dashboard",
        "viewing_camps": "Viewing Campaigns:",
        "active_select": "Active Selection",
        "all_campaigns": "All Campaigns",
        "card_total": "Total Campaigns",
        "card_target": "Total Targets",
        "card_opened": "Opened",
        "card_clicked": "Clicked",
        "card_submitted": "Submitted",
        "card_risk": "Risk Score",
        "risk_critical": "Critical",
        "risk_medium": "Medium",
        "risk_low": "Low",
        "health_ok": "Connected (OK)",
        "health_slow": "Slow",
        "health_err": "Connection Error",
        "no_chart": "No chart yet.\nSync campaigns or make a selection.",
        "dash_no_selection_msg": "Select at least one campaign for a trend.",
        "dash_empty": "No campaigns in storage.\nSync from the 'Campaigns' tab first.",
        "no_campaigns": "No campaigns to display.",
        "report_title_multi": "Multi-Campaign Trend Analysis Report",
        "report_desc_multi": "This report contains a comparison of {n} different campaigns.",
        "report_click_pct": "Clicked (%)",
        "report_submit_pct": "Submitted (%)",
        "report_risk_trend_title": "Risk Trend Over Time",
        "report_oran": "Rate (%)",
        "trend_h_campaign": "Campaign",
        "trend_h_date": "Date",
        "trend_h_target": "Targets",
        "trend_h_clicked": "Clicked (%)",
        "trend_h_submitted": "Submitted (%)",
        "report_title_single": "Corporate Phishing Simulation Report",
        "report_campaign_lbl": "Campaign: {name}",
        "report_exec_summary": "1. Executive Summary",
        "report_pie_critical": "Critical (Submitted)",
        "report_pie_medium": "Medium (Clicked)",
        "report_pie_low": "Low (Opened)",
        "report_pie_ignore": "Non-Engaged",
        "report_stat_target": "Targeted Staff",
        "report_stat_opened": "Total Opened",
        "report_stat_clicked": "Total Clicked",
        "report_stat_submitted": "Submitted (Critical)",
        "report_stat_leak": "Leak (Breach) Rate",
        "report_detail_findings": "2. Detailed Findings",
        "report_table_critical": "Critical Risk: Data-Submitting Users",
        "report_table_medium": "Medium Risk: Users Who Clicked the Link",
        "table_name_col": "Full Name",
        "table_email_col": "Email Address",
        "table_time_col": "Action Time",
        "table_payload_col": "Submitted Data (Payload)",
        "table_device_col": "Device",
        "table_browser_col": "Browser",
        "table_empty": "No users found in this category.",
        "payload_form_empty": "Empty Form",
        "payload_unreadable": "Could not read data",
        "payload_masked": "*** (Hidden)",
        "browser_unavailable": "Device Info Unavailable",
        "excel_title_trend": "Trend Comparison",
        "excel_title_campaign": "Campaign Report",
        "file_report": "_Report",
        "file_trend": "Trend_Analysis",
        "time_all": "All Time",
        "time_1d": "Last 1 Day",
        "time_5d": "Last 5 Days",
        "time_1w": "Last 1 Week",
        "time_1m": "Last 1 Month",
        "selected_msg": "{n} Campaign{s} Selected",
        "drop_hint": "Drag and drop a folder here",
        "w_generate": "Create Word Report",
        "w_add_chart": "Add Chart Analysis",
        "w_add_time": "Add Action Times",
        "w_add_payload": "Add Data (Payload)",
        "w_add_browser": "Add Device/Browser",
        "w_add_dynamic": "Add Form Fields as Separate Columns",
        "e_generate": "Create Excel Report",
        "e_add_time": "Add Action Times",
        "e_add_payload": "Add Submitted Data (Payload)",
        "e_add_browser": "Add Device/Browser Info",
        "e_add_dynamic": "Add Form Fields as Separate Columns",
        "e_add_status": "Add Data-Entry Health Status (Valid/Suspicious/False Positive)",
        "e_add_domain_stats": "Add Domain Breakdown Sheet",
    }
}

def parse_browser_info(event, translations=None):
    try:
        if hasattr(event, 'details') and event.details:
            details = json.loads(event.details) if isinstance(event.details, str) else event.details
            browser = details.get('browser', {})
            os_info = browser.get('os', '')
            user_agent = browser.get('user-agent', '')
            
            if os_info:
                name = browser.get('name', '')
                version = browser.get('version', '')
                return f"{os_info} | {name} {version}".strip(" |")
            elif user_agent:
                return user_agent
    except Exception:
        pass
    return (translations or {}).get("browser_unavailable", "Cihaz Bilgisi Alınamadı")

def _extract_os_ua(ua):
    """User-agent dizesinden isletim sistemi adini/surumunu cikarir."""
    ua = ua or ""
    m = re.search(r"Windows NT (\d+\.\d+)", ua)
    if m:
        ver = m.group(1)
        try:
            major = int(float(ver))
        except Exception:
            major = round(float(ver))
        return "Windows (OS Version: {})".format(major)
    if "iPhone" in ua:
        m = re.search(r"iPhone OS (\d+[\d_]*[\d]*)", ua)
        if m:
            return "Apple iPhone (OS Version: {})".format(m.group(1).replace("_", "."))
        return "Apple iPhone"
    if "iPad" in ua:
        m = re.search(r"iPhone OS (\d+[\d_]*[\d]*)", ua)
        if m:
            return "Apple iPad (OS Version: {})".format(m.group(1).replace("_", "."))
        return "Apple iPad"
    if "Mac OS X" in ua:
        m = re.search(r"Mac OS X (\d+[\d_]*[\d]*)", ua)
        if m:
            return "Mac (OS Version: {})".format(m.group(1).replace("_", "."))
        return "Mac"
    m = re.search(r"Android (\d+(?:\.\d+)*)", ua)
    if m:
        return "Android (OS Version: {})".format(m.group(1))
    if "Linux" in ua or "X11" in ua:
        return "Linux"
    return ""

def _extract_browser_ua(ua):
    """User-agent dizesinden tarayici adi ve surumunu cikarir."""
    ua = ua or ""
    m = re.search(r"Edg(?:A|iOS)?/([\d.]+)", ua)
    if m:
        return "Edge (Version: {})".format(m.group(1))
    m = re.search(r"OPR/([\d.]+)", ua)
    if m:
        return "Opera (Version: {})".format(m.group(1))
    m = re.search(r"Firefox/([\d.]+)", ua)
    if m:
        return "Firefox (Version: {})".format(m.group(1))
    m = re.search(r"Chrome/([\d.]+)", ua)
    if m:
        return "Chrome (Version: {})".format(m.group(1))
    m = re.search(r"Version/([\d.]+).*?Safari/", ua)
    if m:
        return "Safari (Version: {})".format(m.group(1))
    return ""

def parse_device_browser(event, translations=None):
    """Cihaz (OS) ve tarayici bilgisini ayri ayri dondurur: (device, browser)"""
    try:
        if not (hasattr(event, 'details') and event.details):
            return ("-", "-")
        details = json.loads(event.details) if isinstance(event.details, str) else event.details

        browser = details.get('browser', {}) or {}
        if isinstance(browser, str):
            # Browser alani dict degil, dogrudan user-agent dizesi olarak gelmis
            user_agent = browser.strip()
            os_info = ""
            name = ""
            version = ""
        else:
            os_info = str(browser.get('os', '') or '').strip()
            name = str(browser.get('name', '') or '').strip()
            version = str(browser.get('version', '') or '').strip()
            user_agent = str(browser.get('user-agent', '') or '').strip()

        device = os_info
        bt = (name + " (Version: " + version + ")") if (name and version) else name

        if not device and user_agent:
            device = _extract_os_ua(user_agent)
        if not bt and user_agent:
            bt = _extract_browser_ua(user_agent)

        # Hala bulunamazsa, tum details icinde user-agent benzeri herhangi bir alani ara
        if (not bt or not device) and isinstance(details, dict):
            for _k, _v in details.items():
                if isinstance(_v, dict):
                    _ua = str(_v.get('user-agent') or _v.get('user_agent') or '')
                    if _ua:
                        if not bt:
                            bt = _extract_browser_ua(_ua)
                        if not device:
                            device = _extract_os_ua(_ua)
                    if bt and device:
                        break

        return ((device or "-"), (bt or "-"))
    except Exception:
        return ("-", "-")

def format_gophish_date(date_val):
    if not date_val:
        return "-"
    if isinstance(date_val, str):
        try:
            clean_str = date_val.split('.')[0].replace("Z", "")
            dt = datetime.strptime(clean_str, "%Y-%m-%dT%H:%M:%S")
            return dt.strftime("%d.%m.%Y")
        except Exception:
            return date_val[:10]
    else:
        try:
            return date_val.strftime("%d.%m.%Y")
        except Exception:
            return "-"

def get_action_times(timeline, mask_payload=False, translations=None):
    user_times = {}
    sorted_timeline = sorted(timeline, key=lambda x: x.time)

    for event in sorted_timeline:
        email = event.email
        if not email: continue

        if email not in user_times:
            user_times[email] = {
                'Opened': None, 'Clicked': None, 'Submitted': None,
                'Submissions': [], 'RawPayload': "", 'Devices': "", 'Browsers': "", 'Dynamic': {}
            }

        try:
            dt_obj = datetime.strptime(event.time.split('.')[0], "%Y-%m-%dT%H:%M:%S")
            formatted_time = dt_obj.strftime("%d.%m.%Y %H:%M:%S")
        except Exception:
            formatted_time = event.time

        if event.message == 'Email Opened' and not user_times[email]['Opened']:
            user_times[email]['Opened'] = formatted_time
        elif event.message == 'Clicked Link' and not user_times[email]['Clicked']:
            user_times[email]['Clicked'] = formatted_time
        elif event.message == 'Submitted Data':
            if not user_times[email]['Submitted']:
                user_times[email]['Submitted'] = formatted_time

            (device, browser) = parse_device_browser(event, translations)
            payload_str = ""

            try:
                if hasattr(event, 'details') and event.details:
                    details_data = json.loads(event.details) if isinstance(event.details, str) else event.details
                    if 'payload' in details_data:
                        payload = details_data['payload']
                        extracted_fields = []
                        _fields = {}
                        sensitive = ("password", "parola", "sifre", "token", "secret", "api_key", "apikey")
                        for key, val in payload.items():
                            raw = str(val[0]) if isinstance(val, list) and len(val) > 0 else str(val)
                            clean_val = "***" if mask_payload else raw
                            extracted_fields.append(f"{key}: {clean_val}")
                            _val = "***" if (mask_payload and key.lower() in sensitive) else raw
                            _fields[key] = _val
                        payload_str = " | ".join(extracted_fields)
            except Exception:
                payload_str = (translations or {}).get("payload_masked", "*** (Gizlendi)") if mask_payload else (translations or {}).get("payload_unreadable", "Veri okunamadı")

            _fields = {} if '_fields' not in dir() else _fields
            try:
                if hasattr(event, 'details') and event.details:
                    _details_data = json.loads(event.details) if isinstance(event.details, str) else event.details
                    _fields = _details_data.get('payload', {}) or {}
                else:
                    _fields = {}
            except Exception:
                _fields = {}
            _fields = _fields or {}
            sensitive = ("password", "parola", "sifre", "token", "secret", "api_key", "apikey")
            __fields_final = {}
            __fields_raw = {}
            for __k, __v in _fields.items():
                __rv = str(__v[0]) if isinstance(__v, list) and len(__v) > 0 else str(__v)
                __fields_final[__k] = "***" if (mask_payload and __k.lower() in sensitive) else __rv
                __fields_raw[__k] = __rv
            user_times[email]['Submissions'].append({
                'time': formatted_time, 'device': device, 'browser': browser,
                'fields': __fields_final,
                'fields_raw': __fields_raw,
                'payload': payload_str if payload_str else (translations or {}).get('payload_form_empty', 'Form Boş')
            })

    for email, data in user_times.items():
        if data['Clicked'] and not data['Opened']: data['Opened'] = data['Clicked']
        if data['Submitted']:
            if not data['Clicked']: data['Clicked'] = data['Submitted']
            if not data['Opened']: data['Opened'] = data['Submitted']

        submissions = data['Submissions']
        seen_payloads = []
        for s in submissions:
            if s['payload'] not in seen_payloads:
                seen_payloads.append(s['payload'])
        data['RawPayload'] = "\n".join(seen_payloads)
        
        seen_devices = []
        for s in submissions:
            if s['device'] not in seen_devices:
                seen_devices.append(s['device'])
        data['Devices'] = "\n".join(seen_devices)
        seen_browsers = []
        for s in submissions:
            _br = s.get('browser')
            if _br not in seen_browsers:
                seen_browsers.append(_br)
        data['Browsers'] = "\n".join(seen_browsers)

        dyn = {}
        plain_list = []
        for s in submissions:
            _t = s.get('time','')
            _hm = _t.split(" ")[1][:5] if " " in _t else _t
            for k, v in (s.get('fields') or {}).items():
                dyn.setdefault(k, []).append(f"[{_hm}] {v}")
            fr = s.get('fields_raw') or {}
            _p = " | ".join(f"{k}: {v}" for k, v in fr.items())
            if _p and _p not in plain_list:
                plain_list.append(_p)
        data['Dynamic'] = dyn
        data['RawPayloadPlain'] = "\n".join(plain_list)

        raw_dicts = []
        for s in submissions:
            fr = s.get('fields_raw') or {}
            if fr and fr not in raw_dicts:
                raw_dicts.append(fr)
        data['RawPayloadDicts'] = raw_dicts

    return user_times

def set_col_widths(table, dyn_count=0):
    widths = [Inches(1.2), Inches(2.0), Inches(1.4), Inches(2.4), Inches(1.4)]
    for _ in range(dyn_count):
        widths.append(Inches(1.3))
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                break
            try:
                row.cells[idx].width = width
            except IndexError:
                pass


SUSPICIOUS_BLACKLIST = [
    "test", "deneme", "admin", "root", "yok", "fake", "hacked", "hack",
    "parola", "sifre", "password", "123456", "123", "qwe", "asd", "bla",
    "nah", "bos", "none", "null", "bilmiyorum", "gophish"
]

def is_suspicious_payload(raw_payload_dict):
    """Ham payload verisini (maskelenmemiş dict) kontrol eder. Dönüş: (is_suspicious, reason)"""
    if not raw_payload_dict or not isinstance(raw_payload_dict, dict):
        return True, "Bos / Gecersiz Veri"
    for key, val in raw_payload_dict.items():
        val_str = str(val[0] if isinstance(val, list) and len(val) > 0 else val).strip()
        val_lower = val_str.lower()
        if not val_lower:
            continue
        if len(val_str) < 3 and key.lower() in ['password', 'parola', 'sifre']:
            return True, "Cok Kisa Parola ({})".format(val_str)
        for bad_word in SUSPICIOUS_BLACKLIST:
            if bad_word in val_lower:
                return True, "Kara Liste Kelimesi: '{}'".format(bad_word)
        if re.search(r'(.)\1{3,}', val_lower):
            return True, "Karakter Tekrari"
        if re.search(r'(asdf|qwerty|zxcv|12345|67890|qwert|123456)', val_lower):
            return True, "Klavye Deseni"
    return False, "Gecerli"
def add_domain_breakdown_sheet(wb, results, user_times, is_en):
    """Ayri bir calisma sayfasinda domain bazli istatistik tablosu olusturur."""
    domains = {}
    for user in results:
        email = user.email or ""
        if "@" not in email:
            continue
        domain = email.split("@")[-1].lower()
        st = user.status
        d = domains.setdefault(domain, {"total":0,"opened":0,"clicked":0,"submitted":0})
        d["total"] += 1
        if st in ("Email Opened","Clicked Link","Submitted Data"):
            d["opened"] += 1
        if st in ("Clicked Link","Submitted Data"):
            d["clicked"] += 1
        if st == "Submitted Data":
            d["submitted"] += 1

    if not domains:
        return

    sheet_title = "Domain Breakdown" if is_en else "Domain Dağılımı"
    ws = wb.create_sheet(title=sheet_title)

    headers_en = ["DOMAIN","TOTAL TARGETS","OPENED EMAIL","CLICKED LINK","SUBMITTED DATA","CLICK RATE","RISK RATE"]
    headers_tr = ["DOMAİN","TOPLAM HEDEF","E-POSTAYI AÇAN","BAĞLANTIYA TIKLAYAN","VERİ GİRİŞİ YAPAN","TIKLAMA ORANI","RİSK ORANI"]
    headers = headers_en if is_en else headers_tr

    ws.append(headers)
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")

    rows = []
    for dom, mm in domains.items():
        total = mm["total"]
        cr = (mm["clicked"]/total*100) if total else 0
        rr = (mm["submitted"]/total*100) if total else 0
        rows.append((dom, mm["total"], mm["opened"], mm["clicked"], mm["submitted"], round(cr, 1), round(rr, 1)))
    rows.sort(key=lambda r: r[1], reverse=True)

    for r in rows:
        ws.append(list(r))

    t_total = sum(mm["total"] for mm in domains.values())
    t_open  = sum(mm["opened"] for mm in domains.values())
    t_clk   = sum(mm["clicked"] for mm in domains.values())
    t_sub   = sum(mm["submitted"] for mm in domains.values())
    ws.append((("GRAND TOTAL" if is_en else "GENEL TOPLAM"),
               t_total, t_open, t_clk, t_sub,
               round(t_clk/t_total*100, 1) if t_total else 0,
               round(t_sub/t_total*100, 1) if t_total else 0))

    last = ws.max_row
    for cell in ws[last]:
        cell.font = Font(bold=True)

    widths = [32, 16, 18, 20, 22, 15, 15]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[chr(64+i)].width = w
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.freeze_panes = "A2"


def detect_suspicious_payload(payload_str, all_payloads, translations=None):
    pl = (payload_str or "").lower()
    blacklist = ["test","deneme","admin","yok","fake","hacked","parola","sifre","qwerty","asdfgh","123456","12345","11111","111111","asdasd","blabla","anonim","gizli","banka","sikayet","troll"]
    for w in blacklist:
        if w in pl:
            return "troll"
    for kv in pl.split("|"):
        _k, _sep, v = kv.partition(":")
        v = v.strip()
        if 0 < len(v) < 3:
            return "suspicious"
    if re.search(r'(.{2,}?)\1{2,}', pl):
        return "suspicious"
    if all_payloads:
        same = sum(1 for x in all_payloads if (x or "").strip().lower() == pl.strip())
        if same >= 2:
            return "suspicious"
    return "valid"

def status_label(status, translations=None):
    is_en = bool(translations and translations.get("report_oran") == "Rate (%)")
    if status == "troll":
        return "🔴 False Positive / Güvensiz Şifre" if not is_en else "🔴 False Positive / Unsafe Password"
    if status == "suspicious":
        return "🟠 Şüpheli Giriş" if not is_en else "🟠 Suspicious"
    return "🟢 Geçerli" if not is_en else "🟢 Valid"

def create_table(doc, title, data_list, user_times, inc_time, inc_payload, inc_browser, action_key, color_rgb, translations=None, status_map=None, dyn_keys=None):
    has_status = status_map is not None
    has_dyn = bool(dyn_keys) and action_key == 'Submitted'
    heading = doc.add_heading('', level=2)
    run = heading.add_run(title)
    run.font.color.rgb = color_rgb

    if not data_list:
        p = doc.add_paragraph((translations or {}).get("table_empty", "Bu kategoride herhangi bir kullanıcı bulunmamaktadır."))
        p.runs[0].italic = True
        return

    col_count = 2
    if has_status: col_count += 1
    if inc_time: col_count += 1
    if inc_payload and action_key == 'Submitted': col_count += 1
    if inc_browser and action_key == 'Submitted': col_count += 2
    if has_dyn: col_count += len(dyn_keys)

    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    table.autofit = False 

    hdr_cells = table.rows[0].cells
    ci = 0
    if has_status:
        hdr_cells[ci].text = (translations or {}).get('table_status_col', 'Durum'); ci += 1
    hdr_cells[ci].text = (translations or {}).get('table_name_col', 'İsim Soyisim'); ci += 1
    hdr_cells[ci].text = (translations or {}).get('table_email_col', 'E-posta Adresi'); ci += 1
    if inc_time:
        hdr_cells[ci].text = (translations or {}).get('table_time_col', 'Eylem Saati'); ci += 1
    if inc_payload and action_key == 'Submitted':
        hdr_cells[ci].text = (translations or {}).get('table_payload_col', 'Girilen Veri (Payload)'); ci += 1
    if inc_browser and action_key == 'Submitted':
        hdr_cells[ci].text = (translations or {}).get('table_device_col', 'Cihaz'); ci += 1
        hdr_cells[ci].text = (translations or {}).get('table_browser_col', 'Tarayıcı'); ci += 1
    if has_dyn:
        for k in dyn_keys:
            hdr_cells[ci].text = k.upper(); ci += 1

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    for user in data_list:
        row_cells = table.add_row().cells
        full_name = f"{user.first_name} {user.last_name}".strip()
        user_t = user_times.get(user.email, {})
        ci = 0
        if has_status:
            row_cells[ci].text = status_label(status_map.get(user.email, 'valid'), translations); ci += 1
        row_cells[ci].text = full_name if full_name else ("Unknown" if translations and translations.get("report_oran") == "Rate (%)" else "Bilinmiyor"); ci += 1
        row_cells[ci].text = user.email; ci += 1
        if inc_time:
            row_cells[ci].text = user_t.get(action_key, "-") or "-"; ci += 1
        if inc_payload and action_key == 'Submitted':
            row_cells[ci].text = user_t.get('RawPayload', ""); ci += 1
        if inc_browser and action_key == 'Submitted':
            row_cells[ci].text = user_t.get('Devices', ""); ci += 1
            row_cells[ci].text = user_t.get('Browsers', ""); ci += 1
        if has_dyn:
            dyn = user_t.get('Dynamic', {})
            for k in dyn_keys:
                vals = dyn.get(k) or ['-']
                row_cells[ci].text = "\n".join(vals); ci += 1

    set_col_widths(table, dyn_count=len(dyn_keys) if has_dyn else 0)
def sanitize_filename(name):
    if not name: return "Kampanya"
    clean_name = re.sub(r'[\\/*?:"<>|]', '', name).strip(' .')
    return re.sub(r'\s+', '_', clean_name) if clean_name else "Kampanya"

def get_safe_filename(output_dir, base_name, ext):
    final_path = os.path.join(output_dir, f"{base_name}{ext}")
    counter = 1
    while os.path.exists(final_path):
        final_path = os.path.join(output_dir, f"{base_name}_({counter}){ext}")
        counter += 1
    return final_path

class _MergedCampaign:
    def __init__(self, name, timeline, results, created_date):
        self.name=name
        self.timeline=timeline
        self.results=results
        self.created_date=created_date

def merge_campaign_results(campaigns):
    timeline=[]
    results=[]
    for c in campaigns:
        timeline.extend(getattr(c, 'timeline', []) or [])
        results.extend(getattr(c, 'results', []) or [])
    prio={'Submitted Data':4,'Clicked Link':3,'Email Opened':2,'Email Sent':1}
    merged={}
    for r in results:
        em=getattr(r, 'email', None)
        if not em: continue
        if em in merged:
            if prio.get(r.status,0) > prio.get(merged[em].status,0):
                merged[em]=r
        else:
            merged[em]=r
    if campaigns:
        name=", ".join([c.name for c in campaigns]) if len(campaigns)>1 else campaigns[0].name
        created=max([(getattr(c,'created_date','') or '') for c in campaigns], default='')
    else:
        name="Kampanya"; created=""
    return _MergedCampaign(name=name, timeline=timeline, results=list(merged.values()), created_date=created)

# --- TAM RAPOR ÜRETİM MANTIĞI (WORD & EXCEL) ---
def generate_reports(
    api_key: str,
    server: str,
    camp_ids: list,
    output_dir: str,
    options: dict,
    progress: ProgressReporter,
    cancel_event: threading.Event,
    translations: dict,
) -> tuple[list[str], Optional[str]]:
    """
    Seçili kampanyalar için Word/Excel raporları üretir.

    Returns:
        (oluşturulan_dosya_yolları, hata_mesajı)
    """
    generated_files: list[str] = []
    mask_payload = options.get("mask_payload", False)
    client = GophishApiClient(api_key, server, cancel_event)

    progress.advance(translations["step_fetch"], cancel_event)
    try:
        campaigns = []
        for cid in camp_ids:
            ProgressReporter.check_cancel(cancel_event)
            campaigns.append(client.get_campaign(int(cid)))
    except CancelledError:
        raise
    except Exception as e:
        logging.error("API Hatası: %s", e)
        return [], translations["toast_conn_err"].format(msg=str(e))

    os.makedirs(output_dir, exist_ok=True)
    progress.advance(translations["step_process"], cancel_event)

        # --- BİRLEŞİK KAMPANYA RAPORU (1+ KAMPANYA, E-POSTA BAZLI TEKİLLEŞTİRME) ---
    campaign = merge_campaign_results(campaigns)
    user_times = get_action_times(campaign.timeline, mask_payload, translations)

    submitted_data, clicked_link, opened_email = [], [], []
    for result in campaign.results:
        st = result.status
        if st == "Submitted Data":
            submitted_data.append(result); clicked_link.append(result); opened_email.append(result)
        elif st == "Clicked Link":
            clicked_link.append(result); opened_email.append(result)
        elif st == "Email Opened":
            opened_email.append(result)

    # Şüpheli / Troll veri girişi tespiti
    status_map = {}
    for _r in submitted_data:
        _dicts = user_times.get(_r.email, {}).get('RawPayloadDicts', [])
        _st = "valid"
        for _d in _dicts:
            _susp, _reason = is_suspicious_payload(_d)
            if _susp:
                _st = "troll" if "Kara Liste" in _reason else "suspicious"
                break
        status_map[_r.email] = _st

    w_dyn = options.get("w_dynamic_payload", False)
    e_dyn = options.get("e_dynamic_payload", False)
    unique_payload_keys = []
    if w_dyn or e_dyn:
        for _r in submitted_data:
            for _k in user_times.get(_r.email, {}).get('Dynamic', {}):
                if _k not in unique_payload_keys:
                    unique_payload_keys.append(_k)

    is_multi = len(campaigns) > 1
    campaign_names = campaign.name
    if is_multi:
        safe_name = f"Birlestirilmis_Kurumsal_Rapor_{datetime.now().strftime('%Y%m%d_%H%M')}"
    else:
        safe_name = sanitize_filename(campaign.name)

    if options.get("word_enabled"):
        progress.advance(translations["step_word"], cancel_event)

        doc = Document()
        section = doc.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = 1
        section.page_width = new_width
        section.page_height = new_height

        doc.add_heading(translations['report_title_single'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(translations['report_campaign_lbl'].format(name=campaign_names)).bold = True
        doc.add_page_break()

        doc.add_heading(translations['report_exec_summary'], level=1)

        if options.get("w_chart") and len(campaign.results) > 0:
            t = len(campaign.results)
            c_sub = len(submitted_data)
            c_clk = len(clicked_link) - c_sub
            c_opn = len(opened_email) - len(clicked_link)
            c_ign = t - len(opened_email)

            sizes, colors, labels = [], [], []
            for s, c, l in zip([c_sub, c_clk, c_opn, c_ign], ['#e74c3c', '#f39c12', '#f1c40f', '#bdc3c7'], [translations['report_pie_critical'], translations['report_pie_medium'], translations['report_pie_low'], translations['report_pie_ignore']]):
                if s > 0:
                    sizes.append(s); colors.append(c); labels.append(l)

            fig, ax = plt.subplots(figsize=(6, 4))
            ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
            ax.axis('equal')

            memfile = io.BytesIO()
            plt.savefig(memfile, format='png', bbox_inches='tight')
            plt.close(fig)
            memfile.seek(0)
            doc.add_picture(memfile, width=Pt(350))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        stat_table = doc.add_table(rows=5, cols=2)
        stat_table.style = 'Light Shading Accent 1'
        success_rate = (len(submitted_data) / len(campaign.results)) * 100 if len(campaign.results) > 0 else 0

        stats = [
            (translations['report_stat_target'], str(len(campaign.results))),
            (translations['report_stat_opened'], str(len(opened_email))),
            (translations['report_stat_clicked'], str(len(clicked_link))),
            (translations['report_stat_submitted'], str(len(submitted_data))),
            (translations['report_stat_leak'], f"%{success_rate:.2f}".replace('.', ','))
        ]
        for i, (label, value) in enumerate(stats):
            stat_table.rows[i].cells[0].text = label
            stat_table.rows[i].cells[1].text = value
            stat_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_page_break()
        doc.add_heading(translations['report_detail_findings'], level=1)
        create_table(doc, translations['report_table_critical'], submitted_data, user_times, options['w_time'], options['w_payload'], options['w_browser'], 'Submitted', RGBColor(231, 76, 60), translations, status_map, (unique_payload_keys if w_dyn else None))
        doc.add_paragraph("\n")
        create_table(doc, translations['report_table_medium'], clicked_link, user_times, options['w_time'], False, False, 'Clicked', RGBColor(243, 156, 18), translations)

        w_out = get_safe_filename(output_dir, f"{safe_name}{translations['file_report']}", ".docx")
        try:
            doc.save(w_out)
            generated_files.append(w_out)
        except PermissionError:
            return generated_files, translations["toast_perm_err"].format(path=w_out)

    if options.get("excel_enabled"):
        progress.advance(translations["step_excel"], cancel_event)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = translations['excel_title_campaign']

        list_sent = []
        list_opened = []
        list_opened_time = []
        list_clicked = []
        list_clicked_time = []
        list_submitted = []
        list_submitted_time = []
        list_payload = []
        list_device = []
        list_browser = []
        list_status = []

        for user in campaign.results:
            email = user.email
            st = user.status
            u_t = user_times.get(email, {})

            if st in ["Email Sent", "Email Opened", "Clicked Link", "Submitted Data"]:
                list_sent.append(email)

            if st in ["Email Opened", "Clicked Link", "Submitted Data"]:
                list_opened.append(email)
                if options['e_time']: list_opened_time.append(u_t.get('Opened', '-'))

            if st in ["Clicked Link", "Submitted Data"]:
                list_clicked.append(email)
                if options['e_time']: list_clicked_time.append(u_t.get('Clicked', '-'))

            if st == "Submitted Data":
                list_submitted.append(email)
                if options.get('e_status', False): list_status.append(status_label(status_map.get(email, 'valid'), translations))
                if options['e_time']: list_submitted_time.append(u_t.get('Submitted', '-'))
                if options['e_payload']: list_payload.append(u_t.get('RawPayload', '-'))
                if options['e_browser']:
                    list_device.append(u_t.get('Devices', '-'))
                    list_browser.append(u_t.get('Browsers', '-'))

        is_en = (translations['report_oran'] == 'Rate (%)')
        if is_en:
            d_headers = ['USERS EMAILED']; columns_data = [list_sent]
            d_headers.append('USERS OPENED'); columns_data.append(list_opened)
            if options['e_time']: d_headers.append('OPEN TIME'); columns_data.append(list_opened_time)
            d_headers.append('USERS CLICKED'); columns_data.append(list_clicked)
            if options['e_time']: d_headers.append('CLICK TIME'); columns_data.append(list_clicked_time)
            d_headers.append('USERS SUBMITTED DATA'); columns_data.append(list_submitted)
            if options.get('e_status', False): d_headers.append('STATUS / REVIEW'); columns_data.append(list_status)
            if options['e_time']: d_headers.append('SUBMIT TIME'); columns_data.append(list_submitted_time)
            if options['e_payload']: d_headers.append('SUBMITTED DATA (PAYLOAD)'); columns_data.append(list_payload)
            if options['e_browser']:
                d_headers.append('DEVICE'); columns_data.append(list_device)
                d_headers.append('BROWSER'); columns_data.append(list_browser)
            if options.get('e_dynamic_payload', False) and unique_payload_keys:
                for _dk in unique_payload_keys:
                    _dvals = []
                    for _em in list_submitted:
                        _dv = user_times.get(_em, {}).get('Dynamic', {}).get(_dk, ['-'])
                        _dvals.append("\n".join(_dv))
                    d_headers.append(_dk.upper()); columns_data.append(_dvals)
        else:
            d_headers = ['E-POSTA GÖNDERİLEN KULLANICILAR']; columns_data = [list_sent]
            d_headers.append('E-POSTAYI OKUYAN KULLANICILAR'); columns_data.append(list_opened)
            if options['e_time']: d_headers.append('OKUMA SAATİ'); columns_data.append(list_opened_time)
            d_headers.append('BUTONA TIKLAYAN KULLANICILAR'); columns_data.append(list_clicked)
            if options['e_time']: d_headers.append('TIKLAMA SAATİ'); columns_data.append(list_clicked_time)
            d_headers.append('VERİ GİRİŞİ YAPAN KULLANICILAR'); columns_data.append(list_submitted)
            if options.get('e_status', False): d_headers.append('DURUM / İNCELEME'); columns_data.append(list_status)
            if options['e_time']: d_headers.append('VERİ GİRİŞ SAATİ'); columns_data.append(list_submitted_time)
            if options['e_payload']: d_headers.append('GİRİLEN VERİ (PAYLOAD)'); columns_data.append(list_payload)
            if options['e_browser']:
                d_headers.append('CİHAZ'); columns_data.append(list_device)
                d_headers.append('TARAYICI'); columns_data.append(list_browser)
            if options.get('e_dynamic_payload', False) and unique_payload_keys:
                for _dk in unique_payload_keys:
                    _dvals = []
                    for _em in list_submitted:
                        _dv = user_times.get(_em, {}).get('Dynamic', {}).get(_dk, ['-'])
                        _dvals.append("\n".join(_dv))
                    d_headers.append(_dk.upper()); columns_data.append(_dvals)

        ws.append(d_headers)

        green_fill = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
        for cell in ws[1]:
            cell.fill = green_fill
            cell.font = Font(color="000000", bold=True)
            cell.alignment = Alignment(horizontal="left", vertical="center")

        max_len = max([len(c) for c in columns_data]) if columns_data else 0
        for c in columns_data:
            c.extend([''] * (max_len - len(c)))

        for i in range(max_len):
            ws.append([c[i] for c in columns_data])

        for idx, col in enumerate(ws.columns):
            header_text = d_headers[idx]
            if ("SAATİ" in header_text) or ("TIME" in header_text):
                ws.column_dimensions[col[0].column_letter].width = 20
            elif ("VERİ" in header_text and "YAPAN" not in header_text) or ("PAYLOAD" in header_text and "SUBMITTED" in header_text):
                ws.column_dimensions[col[0].column_letter].width = 60
            elif "BROWSER" in header_text or "TARAYICI" in header_text:
                ws.column_dimensions[col[0].column_letter].width = 55
            else:
                ws.column_dimensions[col[0].column_letter].width = 35

        if options.get("e_domain_stats", False):
            add_domain_breakdown_sheet(wb, campaign.results, user_times, is_en)

        e_out = get_safe_filename(output_dir, f"{safe_name}{translations['file_report']}", ".xlsx")
        try:
            wb.save(e_out)
            generated_files.append(e_out)
        except PermissionError:
            return generated_files, translations["toast_perm_err"].format(path=e_out)

    return generated_files, None

# --- MODERN KULLANICI ARAYÜZÜ ---
class GoPhishApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.config_data = load_config()

        ctk.set_appearance_mode(self.config_data.get("theme", "System"))
        self.current_lang = self.config_data.get("lang", "TR")

        self.title(TRANSLATIONS[self.current_lang]["app_title"])
        self.geometry("1100x750")
        self.minsize(950, 650)

        self.all_campaigns = []
        self.selected_campaign_ids = set()

        # Faz 1: Threading altyapısı
        self.ui_queue: queue.Queue = queue.Queue()
        self._cancel_event = threading.Event()
        self._worker_thread: Optional[threading.Thread] = None
        self._is_busy = False

        self.toast = ToastManager(self)

        # Faz 3: Cache ve API saglik izleme
        self.cache = AppCache()
        self.api_health = ApiHealth(self._cancel_event)
        self.last_sync_time: Optional[str] = None

        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        self.init_sidebar()
        self.init_content_frames()
        self.show_frame("dashboard")  # Faz 3: Dashboard ile basla

        self.after(100, self._process_ui_queue)

    def t(self, key):
        return TRANSLATIONS[self.current_lang].get(key, key)

    def _process_ui_queue(self) -> None:
        """Worker thread'lerden gelen mesajları ana thread'de işler."""
        try:
            while True:
                msg = self.ui_queue.get_nowait()
                self._handle_ui_message(msg)
        except queue.Empty:
            pass
        self.after(100, self._process_ui_queue)

    def _handle_ui_message(self, msg: tuple) -> None:
        kind = msg[0]

        if kind == "progress":
            _, step, total, text, percent = msg
            self.progress.set(percent / 100.0)
            self.lbl_status.configure(text=f"{step}/{total} {text}")

        elif kind == "toast":
            self.toast.show(msg[1], msg[2])

        elif kind == "campaigns_loaded":
            _, camps, srv, apikey = msg
            self.all_campaigns = sorted(camps, key=lambda x: x.id, reverse=True)
            self.config_data["server"] = srv
            self.config_data["api_key"] = apikey
            save_config_all(self.config_data)
            self.populate_cards()
            self.last_sync_time = datetime.now().strftime("%H:%M:%S")
            if hasattr(self, "dash_sync_lbl"):
                self.dash_sync_lbl.configure(text=self.last_sync_time)
            if hasattr(self, "dash_camp_combo"):
                names = ["Aktif Secimler", "Tum Kampanyalar"]
                self.dash_camp_combo.configure(values=["Aktif Secimler", "Tum Kampanyalar"])
                self.dash_camp_combo.set("Aktif Secimler")
            self.toast.show(
                self.t("toast_sync_ok").format(count=len(self.all_campaigns)), "success"
            )
            # Faz 3: Cache'e kampanya listesi ozetini yaz
            try:
                brief = [{"id": c.id, "name": c.name, "status": c.status, "created": getattr(c, "created_date", None)} for c in self.all_campaigns]
                self.cache.set("campaign_list", brief)
                self.cache.set("campaign_brief", brief)
            except Exception as exc:
                logging.debug("Cache yazma hatasi: %s", exc)

        elif kind == "campaigns_refreshed":
            _, camps = msg
            self.all_campaigns = sorted(camps, key=lambda x: x.id, reverse=True)
            self.populate_cards()

        elif kind == "report_done":
            _, files, error = msg
            self._reset_busy_state()
            if error:
                self.toast.show(error, "error")
            elif files:
                self.toast.show(
                    self.t("toast_report_ok").format(count=len(files)), "success"
                )
                self.lbl_status.configure(text=self.t("step_done"))
                self.progress.set(1.0)
            else:
                self.lbl_status.configure(text="")
                self.progress.set(0)

        elif kind == "report_cancelled":
            self._reset_busy_state()
            self.lbl_status.configure(text=self.t("step_cancelled"))
            self.progress.set(0)
            self.toast.show(self.t("step_cancelled"), "warning")

        elif kind == "health_result":
            _, data = msg
            self.api_health.status = data.get("status", "unknown")
            self.api_health.latency_ms = data.get("latency_ms")
            self.api_health.last_error = data.get("error")
            self.api_health.checked_at = data.get("checked_at")
            self._update_health_ui()
            if self.api_health.status == "ok":
                self.toast.show(f"API baglanti OK - {self.api_health.latency_ms} ms", "success")
            elif self.api_health.status == "slow":
                self.toast.show(f"API baglanti yavas - {self.api_health.latency_ms} ms", "warning")
            else:
                self.toast.show("API baglanti hatasi: " + str(self.api_health.last_error), "error")

        elif kind == "cache_load":
            _, brief = msg
            from types import SimpleNamespace
            self.all_campaigns = []
            for item in brief:
                self.all_campaigns.append(
                    SimpleNamespace(
                        id=item.get("id"),
                        name=item.get("name", "Bilinmeyen Kampanya"),
                        status=item.get("status", ""),
                        results=[],
                        created_date=item.get("created"),
                    )
                )
            self.all_campaigns.sort(key=lambda x: x.id, reverse=True)
            self.populate_cards()

        elif kind == "fetch_done":
            self._reset_fetch_button()

        elif kind == "refresh_done":
            self._is_busy = False
            self.btn_refresh.configure(state="normal")

    def _reset_busy_state(self) -> None:
        self._is_busy = False
        self._cancel_event.clear()
        if hasattr(self, "btn_gen"):
            self.btn_gen.configure(state="normal", text=self.t("create_report_btn"))
        if hasattr(self, "btn_cancel"):
            self.btn_cancel.pack_forget()

    def _reset_fetch_button(self) -> None:
        self._is_busy = False
        if hasattr(self, "fetch_spinner"):
            self.fetch_spinner.stop(self.t("sync_btn"))
        self.btn_fetch.configure(state="normal")

    def _cancel_operation(self) -> None:
        if self._is_busy:
            self._cancel_event.set()

    def init_sidebar(self):
        self.sidebar_frame = ctk.CTkFrame(self, width=220, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(5, weight=1)
        
        self.logo_label = ctk.CTkLabel(self.sidebar_frame, image=IconManager.get("shield", 24), text=" GoPhish Studio", font=AppFonts.get(16, "bold"), compound="left")
        self.logo_label.grid(row=0, column=0, padx=20, pady=25, sticky="w")

        self.btn_dash = ctk.CTkButton(self.sidebar_frame, image=IconManager.get("sync", 18), text="  Dashboard", command=lambda: self.show_frame("dashboard"), anchor="w", fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"), font=AppFonts.get(13))
        self.btn_dash.grid(row=1, column=0, padx=15, pady=8, sticky="ew")
        
        self.btn_camp = ctk.CTkButton(self.sidebar_frame, image=IconManager.get("campaign", 18), text="  " + self.t("nav_campaigns"), command=lambda: self.show_frame("campaigns"), anchor="w", fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"), font=AppFonts.get(13))
        self.btn_camp.grid(row=2, column=0, padx=15, pady=8, sticky="ew")
        
        self.btn_rep = ctk.CTkButton(self.sidebar_frame, image=IconManager.get("report", 18), text="  " + self.t("nav_reports"), command=lambda: self.show_frame("reports"), anchor="w", fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"), font=AppFonts.get(13))
        self.btn_rep.grid(row=3, column=0, padx=15, pady=8, sticky="ew")
        
        self.btn_set = ctk.CTkButton(self.sidebar_frame, image=IconManager.get("settings", 18), text="  " + self.t("nav_settings"), command=lambda: self.show_frame("settings"), anchor="w", fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"), font=AppFonts.get(13))
        self.btn_set.grid(row=4, column=0, padx=15, pady=8, sticky="ew")
        
        self.appearance_mode_menu = ctk.CTkOptionMenu(self.sidebar_frame, values=["System", "Dark", "Light"], command=self.change_appearance_mode, font=AppFonts.get(12))
        self.appearance_mode_menu.grid(row=6, column=0, padx=20, pady=10, sticky="ew")
        self.appearance_mode_menu.set(self.config_data.get("theme", "System"))
        
        self.lang_menu = ctk.CTkOptionMenu(self.sidebar_frame, values=["TR", "EN"], command=self.change_language, font=AppFonts.get(12))
        self.lang_menu.grid(row=7, column=0, padx=20, pady=(0, 25), sticky="ew")
        self.lang_menu.set(self.current_lang)

    def init_content_frames(self):
        self.container = ctk.CTkFrame(self, fg_color="transparent")
        self.container.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)
        
        self.frames = {}
        
        # 1. Dashboard (Faz 3)
        self.frames["dashboard"] = ctk.CTkFrame(self.container, fg_color="transparent")
        self.build_dashboard_tab(self.frames["dashboard"])
        
        # 2. Kampanyalar
        self.frames["campaigns"] = ctk.CTkFrame(self.container, fg_color="transparent")
        self.build_campaigns_tab(self.frames["campaigns"])
        
        # 3. Rapor Tasarımı
        self.frames["reports"] = ctk.CTkFrame(self.container, fg_color="transparent")
        self.build_reports_tab(self.frames["reports"])
        
        # 3. Ayarlar
        self.frames["settings"] = ctk.CTkFrame(self.container, fg_color="transparent")
        self.build_settings_tab(self.frames["settings"])
        
        for frame in self.frames.values():
            frame.grid(row=0, column=0, sticky="nsew")

    def show_frame(self, name):
        frame = self.frames[name]
        frame.tkraise()

    def change_appearance_mode(self, new_theme):
        ctk.set_appearance_mode(new_theme)
        self.config_data["theme"] = new_theme
        save_config_all(self.config_data)

    def change_language(self, new_lang):
        self.current_lang = new_lang
        self.config_data["lang"] = new_lang
        save_config_all(self.config_data)
        messagebox.showinfo("Bilgi", "Dil tercihi kaydedildi. Lütfen uygulamayı yeniden başlatın.")
        self.destroy()
        sys.exit(0)

    def build_dashboard_tab(self, parent):
        # ---- Faz 3: Dashboard (Risk / API Health / Trend) ----

        # API Health & Senkronizasyon paneli
        health_frame = ctk.CTkFrame(parent)
        health_frame.pack(fill="x", pady=(0, 12))
        health_frame.grid_columnconfigure(3, weight=1)

        ctk.CTkLabel(health_frame, text=self.t("api_health"), font=AppFonts.get(13, "bold")).grid(row=0, column=0, padx=(15, 5), pady=12, sticky="w")
        self.dash_health_badge = ctk.CTkLabel(health_frame, text=self.t("not_checked"), width=110, height=26, corner_radius=6, fg_color=("gray70", "gray30"), text_color="white", font=AppFonts.get(12, "bold"))
        self.dash_health_badge.grid(row=0, column=1, padx=5, pady=12, sticky="w")

        ctk.CTkLabel(health_frame, text=self.t("latency"), font=AppFonts.get(13, "bold")).grid(row=0, column=2, padx=(15, 5), pady=12, sticky="w")
        self.dash_latency_lbl = ctk.CTkLabel(health_frame, text="- ms", font=AppFonts.get(13), text_color=("gray30", "gray70"))
        self.dash_latency_lbl.grid(row=0, column=3, padx=5, pady=12, sticky="w")

        ctk.CTkLabel(health_frame, text=self.t("last_sync"), font=AppFonts.get(13, "bold")).grid(row=0, column=4, padx=(15, 5), pady=12, sticky="w")
        self.dash_sync_lbl = ctk.CTkLabel(health_frame, text="-", font=AppFonts.get(13), text_color=("gray30", "gray70"))
        self.dash_sync_lbl.grid(row=0, column=5, padx=5, pady=12, sticky="w")

        self.btn_health_check = ctk.CTkButton(health_frame, text=self.t("check_btn"), command=self.check_api_health, image=IconManager.get("refresh", 16), font=AppFonts.get(12, "bold"), height=30)
        self.btn_health_check.grid(row=0, column=8, padx=15, pady=12, sticky="e")

        # Risk Skoru Kartlari
        risk_frame = ctk.CTkFrame(parent)
        risk_frame.pack(fill="x", pady=(0, 12))
        self._build_risk_cards(risk_frame)

        # Trend Grafigi
        trend_frame = ctk.CTkFrame(parent)
        trend_frame.pack(fill="both", expand=True, pady=(0, 10))

        trend_header = ctk.CTkFrame(trend_frame, fg_color="transparent")
        trend_header.pack(fill="x", padx=10, pady=8)
        ctk.CTkLabel(trend_header, text=self.t("trend_title"), font=AppFonts.get(15, "bold")).pack(side="left", padx=5)

        self.btn_refresh_dash = ctk.CTkButton(trend_header, text=self.t("refresh_dash"), command=self.refresh_dashboard, image=IconManager.get("refresh", 16), font=AppFonts.get(12, "bold"), height=32)
        self.btn_refresh_dash.pack(side="right", padx=5)

        self.dash_chart_container = ctk.CTkFrame(trend_frame, fg_color=("white", "gray23"))
        self.dash_chart_container.pack(fill="both", expand=True, padx=5, pady=5)
        self.dash_chart_canvas = None
        self.dash_chart_lbl = ctk.CTkLabel(self.dash_chart_container, text=self.t("no_chart"), font=AppFonts.get(12), text_color=("gray40", "gray60"))
        self.dash_chart_lbl.pack(expand=True)

        # Kampanya secim listesi (trend icin)
        select_frame = ctk.CTkFrame(parent, fg_color="transparent")
        select_frame.pack(fill="x", pady=(0, 5))
        ctk.CTkLabel(select_frame, text=self.t("viewing_camps"), font=AppFonts.get(13, "bold")).pack(side="left", padx=10)
        self.dash_camp_combo = ctk.CTkComboBox(select_frame, values=[self.t("active_select"), self.t("all_campaigns")], state="readonly", width=200, font=AppFonts.get(12))
        self.dash_camp_combo.set(self.t("active_select"))
        self.dash_camp_combo.pack(side="left", padx=5)
        self.dash_camp_combo.configure(command=lambda *a: self.refresh_dashboard())

    def _build_risk_cards(self, parent):
        cards = [
            (self.t("card_total"), "0", "#3498db"),
            (self.t("card_target"), "0", "#2c3e50"),
            (self.t("card_opened"), "0", "#2ecc71"),
            (self.t("card_clicked"), "0", "#f39c12"),
            (self.t("card_submitted"), "0", "#e74c3c"),
            (self.t("card_risk"), "0%", "#8e44ad"),
        ]
        for i in range(3):
            parent.grid_columnconfigure(i, weight=1)
        self._risk_card_labels = {}
        for idx, (title, val, color) in enumerate(cards):
            f = ctk.CTkFrame(parent, fg_color=("white", "gray22"), corner_radius=8)
            f.grid(row=idx // 3, column=idx % 3, sticky="nsew", padx=6, pady=6)
            ctk.CTkLabel(f, text=title, font=AppFonts.get(11), text_color=("gray40", "gray60")).pack(padx=10, pady=(10, 2))
            lbl = ctk.CTkLabel(f, text=val, font=AppFonts.get(20, "bold"), text_color=color)
            lbl.pack(padx=10, pady=(0, 10))
            self._risk_card_labels[title] = lbl

    def check_api_health(self):
        srv = self.entry_server.get().strip() if hasattr(self, "entry_server") else self.config_data.get("server", "")
        apikey = self.entry_api.get().strip() if hasattr(self, "entry_api") else self.config_data.get("api_key", "")

        self.dash_health_badge.configure(text=("Check..." if self.current_lang == "EN" else "Kontrol Ediliyor..."), fg_color=("#95a5a6", "#34495e"))
        if hasattr(self, "entry_server"):
            self.entry_server.update_idletasks()

        def worker():
            try:
                data = self.api_health.check(srv, apikey)
                self.ui_queue.put(("health_result", data))
            except Exception as e:
                logging.error("API health hatasi: %s", e)
                self.ui_queue.put(("health_result", {"status": "error", "latency_ms": None, "error": str(e)}))

        threading.Thread(target=worker, daemon=True).start()

    def refresh_dashboard(self):
        """Secili kampanyalarin risk metriklerini ve trend grafigini gunceller."""
        if not hasattr(self, "entry_server"):
            return

        # API Health badge guncelle (bilindiyse)
        self._update_health_ui()

        selected = sorted(self.selected_campaign_ids) if self.selected_campaign_ids else []
        mode = self.dash_camp_combo.get() if hasattr(self, "dash_camp_combo") else self.t("active_select")

        if mode == self.t("all_campaigns") or not selected:
            camps = list(self.all_campaigns)
        else:
            camps = [c for c in self.all_campaigns if c.id in selected]

        if not camps:
            tots = len(self.all_campaigns)
            self._set_risk_card(self.t("card_total"), str(tots))
            self._set_risk_card(self.t("card_target"), "-")
            self._set_risk_card(self.t("card_opened"), "-")
            self._set_risk_card(self.t("card_clicked"), "-")
            self._set_risk_card(self.t("card_submitted"), "-")
            self._set_risk_card(self.t("card_risk"), "-")
            self._show_chart_message(self.t("dash_empty"))
            return

        tot, opn, clk, sub = 0, 0, 0, 0
        names, dates, opened_pct, clicked_pct, submitted_pct = [], [], [], [], []
        for c in camps:
            n = len(c.results) or 0
            s = len([r for r in c.results if r.status == "Submitted Data"])
            k = len([r for r in c.results if r.status in ("Clicked Link", "Submitted Data")])
            o = len([r for r in c.results if r.status in ("Email Opened", "Clicked Link", "Submitted Data")])
            tot += n; opn += o; clk += k; sub += s
            names.append(c.name)
            dates.append(format_gophish_date(getattr(c, 'created_date', None)))
            opened_pct.append(round(o / n * 100, 1) if n > 0 else 0)
            clicked_pct.append(round(k / n * 100, 1) if n > 0 else 0)
            submitted_pct.append(round(s / n * 100, 1) if n > 0 else 0)

        risk = compute_risk_score(opn, clk, sub, tot)
        lvl, lvl_color = risk_level(risk)
        lvl_key = {"Kritik": "risk_critical", "Orta": "risk_medium", "Dusuk": "risk_low"}.get(lvl, "risk_medium")
        lvl_tr = self.t(lvl_key)

        self._set_risk_card(self.t("card_total"), str(len(camps)))
        self._set_risk_card(self.t("card_target"), str(tot))
        self._set_risk_card(self.t("card_opened"), str(opn))
        self._set_risk_card(self.t("card_clicked"), str(clk))
        self._set_risk_card(self.t("card_submitted"), str(sub))
        self._set_risk_card(self.t("card_risk"), f"%{risk} ({lvl_tr})")

        # Trend grafigini ciz
        if len(camps) == 1:
            camp = camps[0]
            t = len(camp.results)
            c_sub = len([r for r in camp.results if r.status == "Submitted Data"])
            c_clk = len([r for r in camp.results if r.status in ("Clicked Link", "Submitted Data")]) - c_sub
            c_opn = len([r for r in camp.results if r.status in ("Email Opened", "Clicked Link", "Submitted Data")]) - (len([r for r in camp.results if r.status in ("Clicked Link", "Submitted Data")]))
            c_ign = t - (t - len([r for r in camp.results if r.status in ("Email Opened", "Clicked Link", "Submitted Data")]))
            sizes = [s for s in (c_sub, c_clk, c_opn, c_ign) if s > 0]
            labels = ["Kritik (Veri)", "Orta (Tik)", "Dusuk (Acan)", "Ilgilenmeyen"][:len(sizes)] if False else []
            sizes2, labels2 = [], []
            for s, l in zip([c_sub, c_clk, c_opn, c_ign], ["Kritik (Veri)", "Orta (Tik)", "Dusuk (Acan)", "Ilgilenmeyen"]):
                if s > 0:
                    sizes2.append(s); labels2.append(l)
            self._show_pie_chart(sizes2, labels2, camp.name)
        elif len(camps) >= 2:
            self._show_trend_line(names, dates, opened_pct, clicked_pct, submitted_pct, mode)
        else:
            self._show_chart_message(self.t("dash_no_selection_msg"))

    def _show_pie_chart(self, sizes, labels, title):
        self._clear_chart_area()
        try:
            from matplotlib.figure import Figure
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            fig = Figure(figsize=(6, 4), dpi=90, facecolor="none")
            ax = fig.add_subplot(111)
            colors = ["#e74c3c", "#f39c12", "#f1c40f", "#bdc3c7"]
            if sizes:
                ax.pie(sizes, labels=labels, colors=colors[:len(sizes)], autopct="%1.1f%%", startangle=140)
            ax.axis("equal")
            ax.set_title(title if title else "Kampanya Dagilimi", fontsize=11)
            canvas = FigureCanvasTkAgg(fig, master=self.dash_chart_container)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both", expand=True, padx=8, pady=8)
            self.dash_chart_canvas = canvas
        except Exception as exc:
            logging.warning("Pie chart gosterilemedi: %s", exc)
            self.dash_chart_lbl.configure(text="Grafik gosterilemedi: " + str(exc))

    def _show_trend_line(self, names, dates, opened_pct, clicked_pct, submitted_pct, mode):
        self._clear_chart_area()
        try:
            from matplotlib.figure import Figure
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            fig = Figure(figsize=(8, 4), dpi=90, facecolor="none")
            ax = fig.add_subplot(111)
            x = list(range(len(names)))
            ax.plot(x, opened_pct, marker="o", label="Acan (%)")
            ax.plot(x, clicked_pct, marker="o", label="Tiklayan (%)")
            ax.plot(x, submitted_pct, marker="o", label="Veri Giren (%)")
            ax.set_xticks(x)
            ax.set_xticklabels(names, rotation=30, ha="right", fontsize=7)
            ax.set_ylabel("Oran (%)")
            ax.set_title("Kampanya Risk Trendi" + (" - Tum Kampanyalar" if mode == "Tum Kampanyalar" else ""), fontsize=11)
            ax.legend()
            fig.tight_layout()
            canvas = FigureCanvasTkAgg(fig, master=self.dash_chart_container)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both", expand=True, padx=8, pady=8)
            self.dash_chart_canvas = canvas
        except Exception as exc:
            logging.warning("Trend grafigi gosterilemedi: %s", exc)
            self.dash_chart_lbl.configure(text="Grafik gosterilemedi: " + str(exc))

    def _show_chart_message(self, text):
        """Grafik alaninda metin mesaji gosterir."""
        self._clear_chart_area()
        if not hasattr(self, "dash_chart_lbl") or not self.dash_chart_lbl.winfo_exists():
            self.dash_chart_lbl = ctk.CTkLabel(self.dash_chart_container, text=text, font=AppFonts.get(12), text_color=("gray40", "gray60"))
        else:
            self.dash_chart_lbl.configure(text=text)
        self.dash_chart_lbl.pack(expand=True)

    def _clear_chart_area(self):
        """Dashboard'daki onceki grafigi temizler ve etiket yeniden olusturur."""
        if hasattr(self, "dash_chart_canvas") and self.dash_chart_canvas is not None:
            try:
                self.dash_chart_canvas.get_tk_widget().destroy()
            except Exception:
                pass
            self.dash_chart_canvas = None
        if not hasattr(self, "dash_chart_lbl") or not self.dash_chart_lbl.winfo_exists():
            self.dash_chart_lbl = ctk.CTkLabel(self.dash_chart_container, text="", font=AppFonts.get(12), text_color=("gray40", "gray60"))
        else:
            self.dash_chart_lbl.pack_forget()

    def _set_risk_card(self, title, value):
        if hasattr(self, "_risk_card_labels") and title in self._risk_card_labels:
            self._risk_card_labels[title].configure(text=value)

    def _update_health_ui(self):
        """Son API health sonucunu dashboard badge'lerine isler."""
        st = self.api_health.status
        if st == "ok":
            self.dash_health_badge.configure(text=self.t("health_ok"), fg_color="#27ae60")
        elif st == "slow":
            self.dash_health_badge.configure(text=self.t("health_slow"), fg_color="#f39c12")
        elif st == "error":
            self.dash_health_badge.configure(text=self.t("health_err"), fg_color="#e74c3c")
        else:
            self.dash_health_badge.configure(text=self.t("not_checked"), fg_color=("gray70", "gray30"))
        if self.api_health.latency_ms is not None:
            self.dash_latency_lbl.configure(text=f"{self.api_health.latency_ms} ms")
        if self.last_sync_time:
            self.dash_sync_lbl.configure(text=self.last_sync_time)

    def build_campaigns_tab(self, parent):
        # Bağlantı Paneli
        conn_frame = ctk.CTkFrame(parent)
        conn_frame.pack(fill="x", pady=(0, 10))
        
        ctk.CTkLabel(conn_frame, text=self.t("server_url"), font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, padx=10, pady=10, sticky="w")
        self.entry_server = ctk.CTkEntry(conn_frame, width=220)
        self.entry_server.insert(0, self.config_data.get("server", ""))
        self.entry_server.grid(row=0, column=1, padx=5, pady=10)
        
        ctk.CTkLabel(conn_frame, text=self.t("api_key"), font=ctk.CTkFont(weight="bold")).grid(row=0, column=2, padx=10, pady=10, sticky="w")
        self.entry_api = ctk.CTkEntry(conn_frame, width=220, show="*")
        self.entry_api.insert(0, self.config_data.get("api_key", ""))
        self.entry_api.grid(row=0, column=3, padx=5, pady=10)
        
        self.btn_fetch = ctk.CTkButton(conn_frame, text=self.t("sync_btn"), command=self.fetch_and_populate, image=IconManager.get("sync", 18), font=AppFonts.get(13), compound="left")
        self.btn_fetch.grid(row=0, column=4, padx=15, pady=10)

        # Filtreleme Paneli
        filter_frame = ctk.CTkFrame(parent, fg_color="transparent")
        filter_frame.pack(fill="x", pady=5)
        
        ctk.CTkLabel(filter_frame, text=self.t("search_lbl"), font=ctk.CTkFont(weight="bold")).pack(side="left", padx=5)
        self.search_var = tk.StringVar()
        self.entry_search = ctk.CTkEntry(filter_frame, textvariable=self.search_var, width=180)
        self.entry_search.pack(side="left", padx=5)
        
        self.var_filter_completed = tk.BooleanVar(value=False)
        self.chk_comp = ctk.CTkCheckBox(filter_frame, text=self.t("completed_chk"), variable=self.var_filter_completed)
        self.chk_comp.pack(side="left", padx=10)
        
        ctk.CTkLabel(filter_frame, text=self.t("date_lbl"), font=ctk.CTkFont(weight="bold")).pack(side="left", padx=(10, 5))
        self.combo_time_filter = ctk.CTkComboBox(filter_frame, values=[self.t("time_all"), self.t("time_1d"), self.t("time_5d"), self.t("time_1w"), self.t("time_1m")], state="readonly", width=130)
        self.combo_time_filter.set(self.t("time_all"))
        self.combo_time_filter.pack(side="left", padx=5)

        self.btn_refresh = ctk.CTkButton(filter_frame, text=f" {self.t('refresh_btn')}", command=self.refresh_campaigns, width=90, image=IconManager.get("refresh", 16), font=AppFonts.get(13), compound="left")
        self.btn_refresh.pack(side="right", padx=5)

        # Verimli Kampanya Listesi (Listbox - cok sayida kampanya icin optimizasyon)
        self.campaign_list_frame = ctk.CTkFrame(parent, corner_radius=10, fg_color=("gray92", "gray17"))
        self.campaign_list_frame.pack(fill="both", expand=True, pady=10)
        self.campaign_list_frame.grid_rowconfigure(0, weight=1)
        self.campaign_list_frame.grid_columnconfigure(0, weight=1)

        self.campaign_listbox = tk.Listbox(
            self.campaign_list_frame,
            selectmode=tk.EXTENDED,
            activestyle="none",
            font=("Segoe UI", 11),
            bg="#2b2b2b" if ctk.get_appearance_mode() == "Dark" else "#ffffff",
            fg="#dcdcdc" if ctk.get_appearance_mode() == "Dark" else "#1a1a1a",
            selectbackground="#1f6aa5",
            selectforeground="#ffffff",
            highlightthickness=0,
            borderwidth=0,
        )
        self.campaign_listbox.grid(row=0, column=0, sticky="nsew", padx=(2, 0), pady=2)
        self.campaign_list_scroll = tk.Scrollbar(self.campaign_list_frame, orient="vertical", command=self.campaign_listbox.yview)
        self.campaign_list_scroll.grid(row=0, column=1, sticky="ns")
        self.campaign_listbox.configure(yscrollcommand=self.campaign_list_scroll.set)

        self.campaign_listbox.bind("<<ListboxSelect>>", self.on_campaign_listbox_select)
        self._campaign_index = {}

        self.search_var.trace("w", lambda *args: self.populate_cards())
        self.var_filter_completed.trace("w", lambda *args: self.populate_cards())
        self.combo_time_filter.configure(command=lambda *args: self.populate_cards())

        # Alt Alt Bilgi Paneli (Live Preview + İlerle Butonu)
        bottom_action_frame = ctk.CTkFrame(parent, fg_color="transparent")
        bottom_action_frame.pack(fill="x", pady=5)

        self.lbl_p_info = ctk.CTkLabel(bottom_action_frame, text=self.t("select_camp_msg"), font=ctk.CTkFont(slant="italic"))
        self.lbl_p_info.pack(side="left", padx=10)

        # Canlı Önizleme Kartları (Seçilince Açılır)
        self.preview_cards_container = ctk.CTkFrame(bottom_action_frame, fg_color="transparent")
        
        def create_preview_card(parent, title, color):
            f = ctk.CTkFrame(parent, fg_color=("gray85", "gray20"), corner_radius=6)
            f.pack(side='left', padx=3)
            ctk.CTkLabel(f, text=title, font=ctk.CTkFont(size=10, weight="bold")).pack(padx=8, pady=1)
            val_lbl = ctk.CTkLabel(f, text="-", font=ctk.CTkFont(size=14, weight="bold"), text_color=color)
            val_lbl.pack(padx=8, pady=1)
            return val_lbl

        self.card_total = create_preview_card(self.preview_cards_container, self.t("target_card"), "#3498db")
        self.card_open = create_preview_card(self.preview_cards_container, self.t("open_card"), "#2ecc71")
        self.card_click = create_preview_card(self.preview_cards_container, self.t("click_card"), "#f39c12")
        self.card_submit = create_preview_card(self.preview_cards_container, self.t("submit_card"), "#e74c3c")

        self.btn_proceed = ctk.CTkButton(bottom_action_frame, text=self.t("proceed_btn"), command=lambda: self.show_frame("reports"), fg_color="#1f6aa5", hover_color="#144870", font=AppFonts.get(13, "bold"), height=35)
        self.btn_proceed.pack(side="right", padx=10)
        self.btn_proceed.configure(state="disabled")

    def fetch_and_populate(self):
        srv, apikey = self.entry_server.get().strip(), self.entry_api.get().strip()
        if not srv or not apikey:
            self.toast.show(self.t("warn_credentials"), "warning")
            return

        if self._is_busy:
            return

        self._is_busy = True
        self._cancel_event.clear()
        self.btn_fetch.configure(state="disabled")
        if not hasattr(self, "fetch_spinner"):
            self.fetch_spinner = LoadingSpinner(self.btn_fetch, self)
        self.fetch_spinner.start(self.t("connecting"), self.t("sync_btn"))

        def worker():
            try:
                client = GophishApiClient(apikey, srv, self._cancel_event)
                camps = client.get_campaigns()
                self.ui_queue.put(("campaigns_loaded", camps, srv, apikey))
            except CancelledError:
                self.ui_queue.put(("toast", self.t("step_cancelled"), "warning"))
            except Exception as e:
                logging.error("Senkronizasyon hatası: %s", e)
                cached = self.cache.get("campaign_brief")
                if cached:
                    self.ui_queue.put(("cache_load", cached))
                    self.ui_queue.put(("toast", "Baglanti kurulamadi; offline cacheten gosteriliyor.", "warning"))
                else:
                    self.ui_queue.put(("toast", self.t("toast_conn_err").format(msg=str(e)), "error"))
            finally:
                self.ui_queue.put(("fetch_done",))

        threading.Thread(target=worker, daemon=True).start()

    def refresh_campaigns(self):
        srv, apikey = self.entry_server.get().strip(), self.entry_api.get().strip()
        if not srv or not apikey:
            return
        if self._is_busy:
            return

        self._is_busy = True
        self._cancel_event.clear()
        self.btn_refresh.configure(state="disabled")

        def worker():
            try:
                client = GophishApiClient(apikey, srv, self._cancel_event)
                camps = client.get_campaigns()
                self.ui_queue.put(("campaigns_refreshed", camps))
            except CancelledError:
                pass
            except Exception as e:
                logging.error("Yenileme hatası: %s", e)
                self.ui_queue.put(("toast", self.t("toast_conn_err").format(msg=str(e)), "error"))
            finally:
                self.ui_queue.put(("refresh_done",))

        threading.Thread(target=worker, daemon=True).start()

    def populate_cards(self):
        self.campaign_listbox.delete(0, tk.END)
        self._campaign_index.clear()
        current_lang = self.current_lang

        search_term = self.search_var.get().lower()
        only_completed = self.var_filter_completed.get()
        time_filter = self.combo_time_filter.get()
        now = datetime.now()

        status_map = {
            "Completed": ("Tamamlandı" if current_lang == "TR" else "Completed"),
            "In Progress": ("Aktif" if current_lang == "TR" else "In Progress"),
            "Error": ("Hatalı" if current_lang == "TR" else "Error"),
        }

        idx = 0
        for c in self.all_campaigns:
            c_date_str = format_gophish_date(getattr(c, 'created_date', None))
            c_name_lower = c.name.lower()

            if search_term and search_term not in c_name_lower and search_term not in c_date_str:
                continue

            status_tr = status_map.get(c.status, c.status)
            if only_completed and status_tr != status_map["Completed"]:
                continue

            if time_filter != self.t("time_all"):
                try:
                    raw_date = getattr(c, 'created_date', None)
                    if isinstance(raw_date, str):
                        clean_str = raw_date.split('.')[0].replace("Z", "")
                        c_dt = datetime.strptime(clean_str, "%Y-%m-%dT%H:%M:%S")
                    elif raw_date:
                        c_dt = raw_date
                    else:
                        c_dt = None
                    if c_dt:
                        delta_days = (now - c_dt).days
                        if time_filter == self.t("time_1d") and delta_days > 1: continue
                        if time_filter == self.t("time_5d") and delta_days > 5: continue
                        if time_filter == self.t("time_1w") and delta_days > 7: continue
                        if time_filter == self.t("time_1m") and delta_days > 30: continue
                except Exception:
                    pass

            # Tek satir, performansli liste girisii
            line = f"[{c.id}]  {c.name}   |   {c_date_str}   |   {status_tr}   |   Hedef: {len(c.results)}"
            self.campaign_listbox.insert(tk.END, line)
            self._campaign_index[idx] = c.id
            idx += 1

        if idx == 0:
            self.campaign_listbox.insert(tk.END, self.t("no_campaigns"))

        self._sync_listbox_selection()

    def on_campaign_listbox_select(self, event):
        """Listbox coklu secimindeki kampanyalari secili sete yazar."""
        selected = self.campaign_listbox.curselection()
        self.selected_campaign_ids = {self._campaign_index[i] for i in selected if i in self._campaign_index}
        self._update_selection_counter()

    def _sync_listbox_selection(self):
        """Mevcut secili kampanyalari listbox'ta tekrar vurgular."""
        for i, cid in self._campaign_index.items():
            if cid in self.selected_campaign_ids:
                self.campaign_listbox.selection_set(i)
            else:
                self.campaign_listbox.selection_clear(i)

    def _update_selection_counter(self):
        count = len(self.selected_campaign_ids)
        if count > 0:
            self.lbl_p_info.configure(text="📊 " + self.t("selected_msg").format(n=count, s="s" if (count != 1 and self.current_lang == "EN") else ""))
            self.preview_cards_container.pack(side="left", padx=10)
            self.btn_proceed.configure(state="normal")
            self.update_live_preview_metrics()
        else:
            self.lbl_p_info.configure(text=self.t("select_camp_msg"))
            self.preview_cards_container.pack_forget()
            self.btn_proceed.configure(state="disabled")

    def update_live_preview_metrics(self):
        tot, opn, clk, sub = 0, 0, 0, 0
        for cid in self.selected_campaign_ids:
            camp = next((c for c in self.all_campaigns if c.id == cid), None)
            if camp:
                tot += len(camp.results)
                sub += len([r for r in camp.results if r.status == "Submitted Data"])
                clk += len([r for r in camp.results if r.status in ["Clicked Link", "Submitted Data"]])
                opn += len([r for r in camp.results if r.status in ["Email Opened", "Clicked Link", "Submitted Data"]])
        
        self.card_total.configure(text=str(tot))
        self.card_open.configure(text=str(opn))
        self.card_click.configure(text=str(clk))
        self.card_submit.configure(text=str(sub))

    def build_reports_tab(self, parent):
        # Üst Panel: Geri Dön Butonu ve Başlık
        top_bar = ctk.CTkFrame(parent, fg_color="transparent")
        top_bar.pack(fill="x", pady=(0, 10))

        btn_back = ctk.CTkButton(top_bar, text=self.t("back_btn"), command=lambda: self.show_frame("campaigns"), fg_color=("gray75", "gray30"), hover_color=("gray65", "gray40"), text_color=("gray10", "gray90"), width=120, height=32, font=AppFonts.get(13, "bold"))
        btn_back.pack(side="left")

        # Çıktı Klasörü Seçimi
        out_frame = ctk.CTkFrame(parent)
        out_frame.pack(fill="x", pady=(0, 15))
        
        ctk.CTkLabel(out_frame, text=self.t("output_dir_lbl"), font=ctk.CTkFont(weight="bold")).pack(side="left", padx=10, pady=12)
        self.entry_output = ctk.CTkEntry(out_frame, width=500)
        self.entry_output.insert(0, self.config_data.get("output_dir", ""))
        self.entry_output.pack(side="left", padx=10, pady=12)
        
        def browse_folder():
            folder = filedialog.askdirectory()
            if folder:
                self.entry_output.delete(0, tk.END)
                self.entry_output.insert(0, folder)
                self.config_data["output_dir"] = folder
                save_config_all(self.config_data)

        ctk.CTkButton(out_frame, text=self.t("browse_btn"), command=browse_folder, width=100).pack(side="left", padx=5)

        # Faz 2: Surukle-birak (klasoru giris alanina birak)
        def on_output_drop(folder_path):
            if folder_path and os.path.isdir(folder_path):
                self.entry_output.delete(0, tk.END)
                self.entry_output.insert(0, folder_path)
                self.config_data["output_dir"] = folder_path
                save_config_all(self.config_data)
                self.toast.show("Klasor secildi: " + os.path.basename(folder_path), "success")

        setup_folder_drop(self.entry_output, on_output_drop)

        # OPSEC Güvenlik
        sec_frame = ctk.CTkFrame(parent)
        sec_frame.pack(fill="x", pady=10)
        self.var_mask = ctk.BooleanVar(value=self.config_data.get("mask_payload", True))
        ctk.CTkCheckBox(sec_frame, text=self.t("mask_chk"), variable=self.var_mask, font=ctk.CTkFont(weight="bold"), text_color="#c0392b").pack(anchor="w", padx=15, pady=12)

        # Rapor Formatları (Word & Excel)
        opts_frame = ctk.CTkFrame(parent, fg_color="transparent")
        opts_frame.pack(fill="both", expand=True, pady=10)
        opts_frame.grid_columnconfigure((0, 1), weight=1)

        word_box = ctk.CTkFrame(opts_frame)
        word_box.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        self.var_word = ctk.BooleanVar(value=self.config_data.get("word_enabled", True))
        self.var_w_chart = ctk.BooleanVar(value=self.config_data.get("w_chart", True))
        self.var_w_time = ctk.BooleanVar(value=self.config_data.get("w_time", False))
        self.var_w_payload = ctk.BooleanVar(value=self.config_data.get("w_payload", False))
        self.var_w_browser = ctk.BooleanVar(value=self.config_data.get("w_browser", False))
        self.var_w_dynamic_payload = ctk.BooleanVar(value=self.config_data.get("w_dynamic_payload", False))

        self.chk_word = ctk.CTkCheckBox(word_box, text=self.t("w_generate"), variable=self.var_word, font=ctk.CTkFont(weight="bold"))
        self.chk_word.pack(anchor="w", padx=15, pady=10)
        self.chk_w_chart = ctk.CTkCheckBox(word_box, text=self.t("w_add_chart"), variable=self.var_w_chart)
        self.chk_w_chart.pack(anchor="w", padx=30, pady=5)
        self.chk_w_time = ctk.CTkCheckBox(word_box, text=self.t("w_add_time"), variable=self.var_w_time)
        self.chk_w_time.pack(anchor="w", padx=30, pady=5)
        self.chk_w_payload = ctk.CTkCheckBox(word_box, text=self.t("w_add_payload"), variable=self.var_w_payload)
        self.chk_w_payload.pack(anchor="w", padx=30, pady=5)
        self.chk_w_browser = ctk.CTkCheckBox(word_box, text=self.t("w_add_browser"), variable=self.var_w_browser)
        self.chk_w_browser.pack(anchor="w", padx=30, pady=5)
        self.chk_w_dynamic_payload = ctk.CTkCheckBox(word_box, text=self.t("w_add_dynamic"), variable=self.var_w_dynamic_payload)
        self.chk_w_dynamic_payload.pack(anchor="w", padx=30, pady=5)

        excel_box = ctk.CTkFrame(opts_frame)
        excel_box.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        
        self.var_excel = ctk.BooleanVar(value=self.config_data.get("excel_enabled", True))
        self.var_e_time = ctk.BooleanVar(value=self.config_data.get("e_time", False))
        self.var_e_payload = ctk.BooleanVar(value=self.config_data.get("e_payload", False))
        self.var_e_browser = ctk.BooleanVar(value=self.config_data.get("e_browser", False))
        self.var_e_dynamic_payload = ctk.BooleanVar(value=self.config_data.get("e_dynamic_payload", False))
        self.var_e_status = ctk.BooleanVar(value=self.config_data.get("e_status", False))
        self.var_e_domain_stats = ctk.BooleanVar(value=self.config_data.get("e_domain_stats", False))

        self.chk_excel = ctk.CTkCheckBox(excel_box, text=self.t("e_generate"), variable=self.var_excel, font=ctk.CTkFont(weight="bold"))
        self.chk_excel.pack(anchor="w", padx=15, pady=10)
        self.chk_e_time = ctk.CTkCheckBox(excel_box, text=self.t("e_add_time"), variable=self.var_e_time)
        self.chk_e_time.pack(anchor="w", padx=30, pady=5)
        self.chk_e_payload = ctk.CTkCheckBox(excel_box, text=self.t("e_add_payload"), variable=self.var_e_payload)
        self.chk_e_payload.pack(anchor="w", padx=30, pady=5)
        self.chk_e_browser = ctk.CTkCheckBox(excel_box, text=self.t("e_add_browser"), variable=self.var_e_browser)
        self.chk_e_browser.pack(anchor="w", padx=30, pady=5)
        self.chk_e_dynamic_payload = ctk.CTkCheckBox(excel_box, text=self.t("e_add_dynamic"), variable=self.var_e_dynamic_payload)
        self.chk_e_dynamic_payload.pack(anchor="w", padx=30, pady=5)
        self.chk_e_status = ctk.CTkCheckBox(excel_box, text=self.t("e_add_status"), variable=self.var_e_status)
        self.chk_e_status.pack(anchor="w", padx=30, pady=5)
        self.chk_e_domain_stats = ctk.CTkCheckBox(excel_box, text=self.t("e_add_domain_stats"), variable=self.var_e_domain_stats)
        self.chk_e_domain_stats.pack(anchor="w", padx=30, pady=5)

        # Ana format acik degilse alt opsiyonlar devre disi kalir (trace ile surekli guncellenir)
        self.var_word.trace_add("write", lambda *a: self._update_report_option_states())
        self.var_excel.trace_add("write", lambda *a: self._update_report_option_states())
        self._update_report_option_states()

        # Üretim Butonu ve Progress
        action_frame = ctk.CTkFrame(parent)
        action_frame.pack(fill="x", pady=15)

        progress_row = ctk.CTkFrame(action_frame, fg_color="transparent")
        progress_row.pack(fill="x", padx=10, pady=(5, 0))

        self.lbl_status = ctk.CTkLabel(progress_row, text="", font=ctk.CTkFont(slant="italic"))
        self.lbl_status.pack(side="left", anchor="w")

        self.btn_cancel = ctk.CTkButton(
            progress_row, text=self.t("cancel_btn"), width=90, height=28,
            fg_color="#c0392b", hover_color="#922b21", font=AppFonts.get(12, "bold"),
            command=self._cancel_operation,
        )

        self.progress = ctk.CTkProgressBar(action_frame)
        self.progress.pack(fill="x", padx=10, pady=5)
        self.progress.set(0)

        def on_generate():
            if not self.selected_campaign_ids:
                self.toast.show(self.t("warn_no_campaign"), "warning")
                self.show_frame("campaigns")
                return

            if not self.var_word.get() and not self.var_excel.get():
                self.toast.show(self.t("warn_no_format"), "warning")
                return

            if self._is_busy:
                return

            camp_ids = list(self.selected_campaign_ids)
            opts = {
                "mask_payload": self.var_mask.get(),
                "word_enabled": self.var_word.get(),
                "w_chart": self.var_w_chart.get(),
                "w_time": self.var_w_time.get(),
                "w_payload": self.var_w_payload.get(),
                "w_browser": self.var_w_browser.get(),
                "w_dynamic_payload": self.var_w_dynamic_payload.get(),
                "excel_enabled": self.var_excel.get(),
                "e_time": self.var_e_time.get(),
                "e_payload": self.var_e_payload.get(),
                "e_browser": self.var_e_browser.get(),
                "e_dynamic_payload": self.var_e_dynamic_payload.get(),
                "e_status": self.var_e_status.get(),
                "e_domain_stats": self.var_e_domain_stats.get(),
            }

            self.config_data.update(opts)
            save_config_all(self.config_data)

            total_steps = count_report_steps(camp_ids, opts)
            translations = TRANSLATIONS[self.current_lang]

            self._is_busy = True
            self._cancel_event.clear()
            self.btn_gen.configure(state="disabled", text=self.t("generating"))
            self.btn_cancel.pack(side="right", padx=5)
            self.lbl_status.configure(text=f"0/{total_steps} ...")
            self.progress.set(0)

            def progress_callback(step, total, text, percent):
                self.ui_queue.put(("progress", step, total, text, percent))

            def worker():
                progress = ProgressReporter(total_steps, progress_callback)
                try:
                    files, error = generate_reports(
                        self.entry_api.get().strip(),
                        self.entry_server.get().strip(),
                        camp_ids,
                        self.entry_output.get().strip(),
                        opts,
                        progress,
                        self._cancel_event,
                        translations,
                    )
                    self.ui_queue.put(("report_done", files, error))
                except CancelledError:
                    self.ui_queue.put(("report_cancelled",))
                except Exception as e:
                    logging.error("Rapor hatası: %s", e)
                    self.ui_queue.put(("report_done", [], str(e)))

            threading.Thread(target=worker, daemon=True).start()

        self.btn_gen = ctk.CTkButton(
            action_frame, text=self.t("create_report_btn"), command=on_generate,
            fg_color="#27ae60", hover_color="#219653",
            font=AppFonts.get(14, "bold"), height=40,
        )
        self.btn_gen.pack(fill="x", padx=10, pady=10)

    def _update_report_option_states(self):
        """Word/Excel ana kutusu kapaliysa alt opsiyonlari devre disi birakir."""
        word_on = self.var_word.get()
        excel_on = self.var_excel.get()
        for chk in (self.chk_w_chart, self.chk_w_time, self.chk_w_payload, self.chk_w_browser, self.chk_w_dynamic_payload):
            chk.configure(state="normal" if word_on else "disabled")
        for chk in (self.chk_e_time, self.chk_e_payload, self.chk_e_browser, self.chk_e_dynamic_payload, self.chk_e_status, self.chk_e_domain_stats):
            chk.configure(state="normal" if excel_on else "disabled")

    def build_settings_tab(self, parent):
        set_frame = ctk.CTkFrame(parent)
        set_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        ctk.CTkLabel(set_frame, text=self.t("settings_title"), font=ctk.CTkFont(size=16, weight="bold")).pack(anchor="w", padx=20, pady=20)
        
        ctk.CTkLabel(set_frame, text=self.t("theme_lbl")).pack(anchor="w", padx=20, pady=5)
        theme_opt = ctk.CTkOptionMenu(set_frame, values=["System", "Dark", "Light"], command=self.change_appearance_mode)
        theme_opt.pack(anchor="w", padx=20, pady=5)
        theme_opt.set(self.config_data.get("theme", "System"))
        
        ctk.CTkLabel(set_frame, text=self.t("lang_lbl")).pack(anchor="w", padx=20, pady=(20, 5))
        lang_opt = ctk.CTkOptionMenu(set_frame, values=["TR", "EN"], command=self.change_language)
        lang_opt.pack(anchor="w", padx=20, pady=5)
        lang_opt.set(self.current_lang)

if __name__ == "__main__":
    app = GoPhishApp()
    app.mainloop()